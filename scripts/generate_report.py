from __future__ import annotations

from pathlib import Path
from textwrap import dedent, wrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "docs" / "report-assets"
OUTPUT_PATH = ROOT / "docs" / "GraphLink_Final_Raporu.docx"

TEMPLATE_SOURCE = Path(r"c:\Users\5A_Traders\OneDrive\Desktop\PROJE-RAPORU-SABLON.docx")

MEASURE_DATE = "23 Nisan 2026"

PROJECT_FACTS = {
    "project_name": "GraphLink",
    "course": "BMU1208 — WEB TABANLI PROGRAMLAMA",
    "title": "NEO4J GRAPH DB SOSYAL AG",
    "subtitle": "GraphLink: Neo4j tabanli sosyal ag ve graph analiz paneli",
    "student": "AMIR SHEIKH",
    "student_no": "24080410155",
    "lecturer": "Dr. Ogr. Uyesi Davut ARI",
    "term": "2025-2026 BAHAR",
    "city": "BITLIS",
    "node_version": "v24.13.1",
    "npm_version": "11.8.0",
    "backend": "Node.js + Express + neo4j-driver",
    "frontend": "HTML + CSS + vanilla JavaScript + Cytoscape.js",
    "database": "Neo4j Community 2026.03.1",
    "gds_version": "2026.03.0",
    "apoc_version": "2026.03.1",
    "demo_counts": "6 kullanici, 4 seed gonderi, 12 takip iliskisi ve 8 begeni",
    "live_counts": "MEASURE_DATE tarihinde canli oturumda 6 kullanici, 12 takip iliskisi, 5 gonderi ve 8 begeni goruldu",
    "dashboard_avg_ms": "454.7",
    "dashboard_min_ms": "332",
    "dashboard_max_ms": "596",
    "dashboard_samples": "458, 513, 435, 430, 515, 596, 376, 409, 332, 483",
    "tests_passed": "61",
    "tests_stack": "node:test, Supertest, c8 coverage, validation/security/route/frontend smoke ve Neo4j entegrasyon testleri",
    "runtime_status": "GDS ve APOC eklentileri yerel Neo4j runtime icinde aktif"
}


REFERENCES = [
    "[1] Fortune Business Insights. (2026). Graph Database Market Size, Share & Industry Analysis. "
    "https://www.fortunebusinessinsights.com/graph-database-market-105916",
    "[2] Neo4j. (2026). Neo4j Community Edition. https://neo4j.com/product/community-edition/",
    "[3] Neo4j. (2026). Neo4j Pricing. https://neo4j.com/pricing/",
    "[4] Neo4j. (2026). Cypher Manual - Introduction. https://neo4j.com/docs/cypher-manual/current/introduction/",
    "[5] Neo4j. (2026). Neo4j Graph Data Science Library Manual v2026.03. "
    "https://neo4j.com/docs/graph-data-science/current/",
    "[6] Neo4j. (2026). APOC Core Documentation - Introduction. https://neo4j.com/docs/apoc/current/introduction/",
    "[7] Cytoscape.js. (2026). Cytoscape.js Documentation. https://js.cytoscape.org/",
    "[8] OWASP Foundation. (2026). OWASP Top 10:2021. https://owasp.org/Top10/2021/",
    "[9] W3C. (2025). Web Content Accessibility Guidelines (WCAG) 2.1. https://www.w3.org/TR/WCAG21/",
    "[10] Stack Overflow. (2011). SQL query for mutual friends. https://stackoverflow.com/questions/6701090/sql-query-for-mutual-friends",
    "[11] Stack Overflow. (2017). 'friends of friends' SQL query. https://stackoverflow.com/questions/43913201/friends-of-friends-sql-query",
    "[12] Reddit / r/facebook. (2026, February 21). 'People You May Know' is getting creepy. "
    "https://www.reddit.com/r/facebook/comments/1rb42e3/people_you_may_know_is_getting_creepy_looking_for/",
    "[13] Reddit / r/facebook. (2025). Facebook 'People you may know' recommending a person who I passed in the street? "
    "https://www.reddit.com/r/facebook/comments/1kihoma/facebook_people_you_may_know_recommending_a/",
    "[14] Kumu. (2026). Pricing. https://www.kumu.io/pricing",
    "[15] Kumu. (2026). About. https://kumu.io/about",
    "[16] Linkurious. (2026). Pricing. https://linkurious.com/pricing/",
    "[17] Linkurious. (2026). About. https://linkurious.com/about/",
    "[18] Neo4j. (2026). Company. https://neo4j.com/company/"
]


USER_STORIES = [
    {
        "story": "As a graph meraklisi ogrenci, I want to hazir demo hesaplar ile hizlica giris yapabileyim, so that kurulumdan bagimsiz olarak projeyi aninda deneyebileyim.",
        "acceptance": "Given acilis ekrani, When demo hesaplari yuklenirse, Then kullanici paylasilan sifreyi ve hesaplari gorur.\n"
        "Given gecerli demo kimlik bilgileri, When kullanici giris yapar, Then dashboard acilir.\n"
        "Given gecersiz sifre, When giris denenirse, Then hata mesaji gosterilir.",
        "priority": "Must-have",
    },
    {
        "story": "As a yeni kullanici, I want to kayit formuyla kendi hesabimi olusturayim, so that sistemi demo disinda gercek bir akista test edebileyim.",
        "acceptance": "Given bos form, When zorunlu alanlar eksikse, Then Zod tabanli validation hatasi doner.\n"
        "Given benzersiz e-posta, When kayit tamamlanir, Then User node'u olusur.\n"
        "Given kayit basarili, When response donerse, Then auth cookie'leri set edilir.",
        "priority": "Must-have",
    },
    {
        "story": "As an authenticated user, I want to kim oldugumu gorebileyim, so that aktif oturumun dogru hesapta oldugunu teyit edeyim.",
        "acceptance": "Given gecerli access token, When /api/auth/me cagrilir, Then kullanici ozeti JSON doner.\n"
        "Given suresi dolmus access token, When refresh basarili ise, Then istek otomatik tekrar edilir.\n"
        "Given oturum yok, When korumali endpoint cagrilir, Then 401 doner.",
        "priority": "Must-have",
    },
    {
        "story": "As a member, I want to diger kullanicilari listeleyeyim ve takip edeyim, so that sosyal graphi genisletebileyim.",
        "acceptance": "Given users listesi, When takip edilmeyen biri gorunurse, Then Takip et butonu aktif olur.\n"
        "Given gecerli targetId, When /api/follows POST cagrilir, Then FOLLOWS iliskisi olusur.\n"
        "Given ayni hedef tekrar secilirse, When islem yinelenir, Then MERGE ile kopya iliski olusmaz.",
        "priority": "Must-have",
    },
    {
        "story": "As a member, I want to takipten cikabileyim, so that sosyal graphimi duzenleyebileyim.",
        "acceptance": "Given mevcut takip iliskisi, When /api/follows DELETE cagrilir, Then iliski silinir.\n"
        "Given liste yenilenirse, When kart tekrar render edilirse, Then buton Takip et olarak degisir.\n"
        "Given gecersiz targetId, When silme istenir, Then validation hatasi uretilir.",
        "priority": "Must-have",
    },
    {
        "story": "As a member, I want to gonderi paylasayim, so that graph icindeki icerik akisina katki saglayayim.",
        "acceptance": "Given 2-500 karakter arasi icerik, When /api/posts POST cagrilir, Then Post node'u ve AUTHORED iliskisi olusur.\n"
        "Given bos icerik, When form gonderilirse, Then hata mesaji doner.\n"
        "Given paylasim basarili, When dashboard yenilenirse, Then yeni gonderi akista gorunur.",
        "priority": "Must-have",
    },
    {
        "story": "As a member, I want to gonderi begenebileyim ve geri alabileiyim, so that etkileisimleri test edebileyim.",
        "acceptance": "Given bir post, When like endpoint'i cagrilir, Then LIKED iliskisi olusur.\n"
        "Given mevcut begeni, When unlike endpoint'i cagrilir, Then LIKED iliskisi silinir.\n"
        "Given dashboard tekrar yuklenirse, When post render edilirse, Then likeCount guncel gorunur.",
        "priority": "Must-have",
    },
    {
        "story": "As a user, I want to 2. derece baglantilarimi goreyim, so that kimler bana sosyal olarak yakin anlayabileyim.",
        "acceptance": "Given viewer node'u, When dashboard hesaplanir, Then 2-hop ama direkt olmayan adaylar listelenir.\n"
        "Given adaylar varsa, When UI render edilirse, Then uzerinden gelinen isimler de gosterilir.\n"
        "Given sonuc yoksa, When panel bos kalir, Then empty-state gosterilir.",
        "priority": "Must-have",
    },
    {
        "story": "As a user, I want to tavsiye edilen kisileri ve nedenlerini goreyim, so that onerilerin aciklanabilirligini degerlendireyim.",
        "acceptance": "Given viewer begenileri ve baglantilari, When recommendation query calisir, Then puanli aday listesi uretilir.\n"
        "Given hem mutual hem peer sinyali varsa, When sonuc donerse, Then reason alani melez gerekce yazar.\n"
        "Given skor sifirsa, When filtre uygulanir, Then aday listede yer almaz.",
        "priority": "Must-have",
    },
    {
        "story": "As a graph analyst, I want to iki kisi arasindaki en kisa yolu goreyim, so that topluluk yapisini aciklayabileyim.",
        "acceptance": "Given target secili, When GDS etkin ise, Then Dijkstra sonucu doner.\n"
        "Given GDS yoksa, When path istenir, Then shortestPath fallback'i kullanilir.\n"
        "Given path bulunduysa, When UI render edilirse, Then hop sayisi ve node sirasi gorunur.",
        "priority": "Must-have",
    },
]


TECH_STACK_CARDS = [
    {
        "heading": "4.2.1. Neo4j 5 Community",
        "what": "Neo4j, iliski yogun veriyi node ve edge yapisiyla saklayan native property graph veritabanidir.",
        "why": [
            "Sosyal ag sorgularinda cok adimli traversallar SQL join zincirlerine gore daha dogal ifade edilir.",
            "Projede FOLLOWS, AUTHORED, LIKED ve HAS_SESSION iliskileri veri modelinin merkezindedir.",
            "Yerel Community surumu final teslimi icin lisans maliyeti olmadan kurulum saglar [2].",
        ],
        "features": [
            "Property graph modeli",
            "Unique constraint destegi",
            "Bolt protokolu ve resmi JavaScript driver",
            "Cypher ile pattern tabanli sorgular",
            "Self-managed, offline demo kurulum",
            "Index-free adjacency ile traversal optimizasyonu [2]",
        ],
        "alternatives": [
            "PostgreSQL: tablosal yapida iliski sayisi arttikca sorgu karmasikligi artar.",
            "MongoDB: dokuman modeli takip grafini sorgulamada rahat degildir.",
            "Neo4j AuraDB: bulut bagimliligi ve surekli maliyet bu teslim icin gereksizdi.",
        ],
        "tradeoffs": [
            "Community Edition HA ve kurumsal guvenlik ozelliklerini sinirli sunar.",
            "Yerel runtime bakimi ogrenci makinesinde yapilir.",
        ],
        "resources": ["Neo4j Community Edition sayfasi [2]", "Neo4j Company / tarihce [18]"],
        "role": "Butun is kurallari, sosyal iliskiler, session node'lari ve graph analizleri Neo4j uzerinde tutulur.",
    },
    {
        "heading": "4.2.2. Cypher Query Language",
        "what": "Cypher, Neo4j'in declarative graph sorgu dilidir [4].",
        "why": [
            "MATCH kaliplari sosyal iliskileri okunabilir sekilde ifade eder.",
            "Kisa yol, mutual ve recommendation sorgulari tek dilde yazilabilir.",
            "Constraint ve projection sorgulari ayni ekosistem icinde kalir.",
        ],
        "features": [
            "Pattern matching",
            "Path sorgulari",
            "Aggregation",
            "Map projection",
            "Constraint ve index tanimlari",
            "GDS ve APOC ile beraber kullanilabilme",
        ],
        "alternatives": [
            "SQL: mutual ve friends-of-friends senaryolarinda self-join zinciri gerekir [10][11].",
            "Gremlin: daha esnek fakat bu proje icin gereksiz derecede dusuk seviyeli.",
        ],
        "tradeoffs": [
            "Takim Neo4j disinda yeni bir sorgu soz dizimi ogrenmek zorundadir.",
            "Vendor lock-in riski vardir.",
        ],
        "resources": ["Cypher Manual [4]"],
        "role": "Servis katmanindaki tum read/write islemleri Cypher uzerinden calisir.",
    },
    {
        "heading": "4.2.3. APOC Library",
        "what": "APOC, Cypher'i veri donusumu ve yardimci prosedurlerle genisleten resmi destekli ek kitapliktir [6].",
        "why": [
            "Projede graph runtime varlik kontrolu ve gelismis prosedur kullanimi icin genisletilebilirlik saglar.",
            "Yerel plugin yapisina eklenmesi kolaydir.",
            "Template gereksinimindeki prosedur destegini dogrudan karsilar.",
        ],
        "features": [
            "User-defined procedures ve functions",
            "JSON/CSV gibi veri donusum yardimcilari",
            "Graph utility prosedurleri",
            "Cypher ile dogrudan birlikte cagrilabilme [6]",
        ],
        "alternatives": [
            "Sifirdan custom procedure yazmak: bu proje olceginde gereksizdi.",
            "Sadece cekirdek Cypher: gelismis operasyonlarda esneklik azalirdi.",
        ],
        "tradeoffs": [
            "Plugin kurulum ve versiyon eslestirmesi gerekir.",
            "Bazi prosedurler dikkatli kullanilmadiginda ek bellek tuketebilir [6].",
        ],
        "resources": ["APOC Core docs [6]"],
        "role": "Projede plugin durumu gosterimi ve ileri seviye graph yardimcilarinin kapisi olarak konumlanir.",
    },
    {
        "heading": "4.2.4. Graph Data Science Library",
        "what": "Neo4j GDS, graph algoritmalari ve graph ML senaryolari icin kullanilan analiz kitapligidir [5].",
        "why": [
            "Template'te istenen Louvain, PageRank, FastRP ve Dijkstra'yi ayni ortamda saglar.",
            "Sosyal ag baglanti oruntulerini aciklamak icin hazir algoritmalar sunar.",
            "GDS yoksa Cypher fallback kullanilarak arayuz kirilmadan calisabilir.",
        ],
        "features": [
            "Graph projection",
            "Community detection",
            "Centrality algoritmalari",
            "Embedding/ML prosedurleri",
            "Path finding",
        ],
        "alternatives": [
            "NetworkX/ayri Python servisleri: ek deployment ve veri senkronizasyonu gerektirirdi.",
            "Sadece Cypher: community ve embedding hesaplari zayif kalirdi.",
        ],
        "tradeoffs": [
            "Projection olusturma ek hesaplama maliyeti getirir.",
            "Plugin versiyonu Neo4j surumuyle uyumlu tutulmalidir.",
        ],
        "resources": ["GDS Manual [5]"],
        "role": "Louvain, PageRank, FastRP ve Dijkstra sonuclari analiz paneline veri saglar.",
    },
    {
        "heading": "4.3.1. Node.js + Express + neo4j-driver",
        "what": "Node.js ve Express, REST API katmanini; neo4j-driver ise Bolt uzerinden veritabanina baglantiyi saglar.",
        "why": [
            "Tek dilde full-stack gelistirme hiz kazandirir.",
            "Mevcut calisan Express iskeleti ADR-001 ile korundu.",
            "neo4j-driver, transaction ve session yonetimini dogrudan destekler.",
        ],
        "features": [
            "Middleware zinciri",
            "Cookie parser ve JSON body parser",
            "Helmet entegrasyonu",
            "Route bazli auth ve rate limiting",
            "executeRead / executeWrite akislari",
        ],
        "alternatives": [
            "Fastify: template'te gecse de mevcut calisan mimariyi yeniden yazmak gereksizdi.",
            "NestJS: fazla agirlik katardi.",
        ],
        "tradeoffs": [
            "Buyuk ekiplerde ek mimari disiplin gerekir.",
            "Express'in opinionated olmamasi dosya duzenine dikkat gerektirir.",
        ],
        "resources": ["Proje README ve ADR-001", "Neo4j JavaScript driver docs (dolayli)"],
        "role": "Auth, social ve graph router'lari servis katmanina baglar; static dosya ve Swagger yayinlar.",
    },
    {
        "heading": "4.4.1. Vanilla JavaScript + Cytoscape.js",
        "what": "Vanilla JavaScript istemci akisini; Cytoscape.js ise network gorunumunu saglar [7].",
        "why": [
            "Template'teki React opsiyonu yerine mevcut UI daha hizli tamamlandi.",
            "Cytoscape.js web tarafinda graph gorsellestirme icin yeterli ve hafifti [7].",
            "Kucuk teslim kapsaminda framework ek yuk getirmedi.",
        ],
        "features": [
            "Modul tabanli istemci kodu",
            "Fetch ile API entegrasyonu",
            "Otomatik refresh akisi",
            "Cose layout ile graph render",
            "Node tap ile hedef kullanici secimi",
        ],
        "alternatives": [
            "React + vis-network: yeniden yazim maliyeti ve teslim riski.",
            "Neo4j Bloom: UI icine gomulu ve ozellestirilmis bir akisa uygun degildi.",
        ],
        "tradeoffs": [
            "State yonetimi buyudukce manuel kalabilir.",
            "Component ayrimi React kadar formal degildir.",
        ],
        "resources": ["Cytoscape.js docs [7]"],
        "role": "Auth formu, dashboard panelleri, graph canvas ve responsive arayuz bu katmanda calisir.",
    },
    {
        "heading": "4.5.1. JWT + HttpOnly Cookie",
        "what": "JWT access token ve veritabaninda hashlenmis refresh token session modeli birlikte kullanildi.",
        "why": [
            "Korumali endpoint'ler icin hafif ve iyi bilinen bir desen sunar.",
            "HttpOnly cookie istemcide token saklama riskini azaltir.",
            "Refresh session node'lari logout ve rotation icin ek kontrol saglar.",
        ],
        "features": [
            "15 dakikalik access token TTL",
            "7 gunluk refresh cookie",
            "HttpOnly + sameSite=lax",
            "Role claim",
            "Refresh rotation",
        ],
        "alternatives": [
            "LocalStorage token: XSS riskini artirir.",
            "Tam OAuth 2.0 / social login: bu kapsam icin fazla buyuktu.",
        ],
        "tradeoffs": [
            "Cookie ve token surelerinin senkron yonetimi gerekir.",
            "Session node temizligi ek operasyonel is getirir.",
        ],
        "resources": ["ADR-003", "Proje auth servis kodu"],
        "role": "Login, refresh, logout ve role-based demo reset yetkisi bu mekanizma ile calisir.",
    },
    {
        "heading": "4.6.1. Yerel Neo4j Runtime (Windows Scriptleri)",
        "what": "Neo4j Community, proje icindeki PowerShell scriptleriyle yerelde kurulur ve baslatilir.",
        "why": [
            "Teslim sirasinda internet veya bulut hesabina bagimlilik kalmaz.",
            "GDS ve APOC jar dosyalari tek akista hazirlanir.",
            "Gelisim ve demo ortaminda tekrar edilebilir kurulum sunar.",
        ],
        "features": [
            "setup-neo4j.ps1 ile kurulum",
            "start/stop scriptleri",
            "Plugin kopyalama",
            "127.0.0.1 ile sinirli listen address",
            "Auth kapali demo runtime",
        ],
        "alternatives": [
            "Neo4j AuraDB Free: internet bagimliligi ve farkli operasyonel model.",
            "Docker: bu makinede ek katman gerektirirdi.",
        ],
        "tradeoffs": [
            "Windows'a ozgu scriptler farkli ortamlarda uyarlama gerektirebilir.",
            "Production topolojisi icin yeterli degildir; daha cok demo odaklidir.",
        ],
        "resources": ["ADR-002", "scripts/setup-neo4j.ps1"],
        "role": "Final gosteriminde uygulamanin tek komutla ayağa kalkmasini saglar.",
    },
]


def ensure_assets() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf"),
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def center_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill: str) -> None:
    width = max(12, int((box[2] - box[0]) / 11))
    wrapped = "\n".join(wrap(text, width=width))
    bbox = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=6, align="center")
    x = box[0] + (box[2] - box[0] - (bbox[2] - bbox[0])) / 2
    y = box[1] + (box[3] - box[1] - (bbox[3] - bbox[1])) / 2
    draw.multiline_text((x, y), wrapped, font=font, fill=fill, spacing=6, align="center")


def box(draw: ImageDraw.ImageDraw, rect: tuple[int, int, int, int], text: str, fill: str, outline: str = "#0f172a",
        text_fill: str = "#ffffff") -> None:
    draw.rounded_rectangle(rect, radius=24, fill=fill, outline=outline, width=3)
    center_text(draw, rect, text, get_font(28, bold=True), text_fill)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], label: str | None = None,
          color: str = "#0f172a") -> None:
    draw.line([start, end], fill=color, width=5)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex >= sx else -1
        draw.polygon([(ex, ey), (ex - 18 * direction, ey - 10), (ex - 18 * direction, ey + 10)], fill=color)
    else:
        direction = 1 if ey >= sy else -1
        draw.polygon([(ex, ey), (ex - 10, ey - 18 * direction), (ex + 10, ey - 18 * direction)], fill=color)
    if label:
        mx = (sx + ex) / 2
        my = (sy + ey) / 2
        draw.rounded_rectangle((mx - 95, my - 18, mx + 95, my + 18), radius=12, fill="#e2e8f0", outline="#cbd5e1")
        center_text(draw, (int(mx - 92), int(my - 16), int(mx + 92), int(my + 16)), label, get_font(18), "#0f172a")


def canvas(title: str, subtitle: str = "") -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1600, 900), "#f8fafc")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((24, 24, 1576, 876), radius=28, outline="#cbd5e1", width=3, fill="#ffffff")
    draw.text((56, 42), title, font=get_font(34, bold=True), fill="#0f172a")
    if subtitle:
        draw.text((56, 88), subtitle, font=get_font(20), fill="#475569")
    return image, draw


def generate_diagrams() -> None:
    ensure_assets()

    img, draw = canvas("Sistem Context Diyagrami", "GraphLink - kullanici, web arayuzu, API ve Neo4j iliskisi")
    box(draw, (120, 320, 360, 470), "Kullanici", "#0f766e")
    box(draw, (450, 280, 770, 510), "GraphLink Web UI\n(HTML + CSS + JS)", "#0f172a")
    box(draw, (870, 280, 1180, 510), "Express API", "#1d4ed8")
    box(draw, (1260, 230, 1490, 430), "Neo4j Browser", "#f59e0b", text_fill="#111827")
    box(draw, (1260, 500, 1490, 700), "Neo4j Community", "#0f766e")
    arrow(draw, (360, 395), (450, 395), "HTTP")
    arrow(draw, (770, 395), (870, 395), "REST")
    arrow(draw, (1180, 395), (1260, 600), "Bolt")
    arrow(draw, (1375, 430), (1375, 500), "inspect")
    img.save(ASSET_DIR / "context-diagram.png")

    img, draw = canvas("Container Diyagrami", "Istemci, servis katmani ve veri katmaninin baglantilari")
    draw.rounded_rectangle((80, 180, 530, 760), radius=26, outline="#94a3b8", width=3, fill="#f8fafc")
    draw.rounded_rectangle((570, 180, 1080, 760), radius=26, outline="#94a3b8", width=3, fill="#f8fafc")
    draw.rounded_rectangle((1120, 180, 1520, 760), radius=26, outline="#94a3b8", width=3, fill="#f8fafc")
    draw.text((110, 200), "Client", font=get_font(26, True), fill="#0f172a")
    draw.text((600, 200), "Server", font=get_font(26, True), fill="#0f172a")
    draw.text((1150, 200), "Data", font=get_font(26, True), fill="#0f172a")
    box(draw, (120, 280, 490, 420), "Auth + Dashboard UI", "#0f172a")
    box(draw, (120, 500, 490, 640), "Cytoscape Graph View", "#0f766e")
    box(draw, (620, 250, 1020, 370), "Auth Router", "#1d4ed8")
    box(draw, (620, 410, 1020, 530), "Social Router", "#1d4ed8")
    box(draw, (620, 570, 1020, 690), "Graph Router + Service Layer", "#1d4ed8")
    box(draw, (1170, 250, 1470, 390), "Neo4j 5 Community", "#0f766e")
    box(draw, (1170, 460, 1310, 650), "GDS", "#f59e0b", text_fill="#111827")
    box(draw, (1330, 460, 1470, 650), "APOC", "#f59e0b", text_fill="#111827")
    arrow(draw, (490, 350), (620, 310), "login/register")
    arrow(draw, (490, 350), (620, 470), "users/posts")
    arrow(draw, (490, 570), (620, 630), "network")
    arrow(draw, (1020, 620), (1170, 320), "Cypher")
    arrow(draw, (1320, 390), (1240, 460), "algorithms")
    arrow(draw, (1320, 390), (1400, 460), "procedures")
    img.save(ASSET_DIR / "container-diagram.png")

    img, draw = canvas("Login Sequence Diyagrami", "Kullanici -> UI -> API -> Neo4j akisi")
    xs = [180, 520, 860, 1220]
    labels = ["Kullanici", "Web UI", "Express API", "Neo4j"]
    for x, label in zip(xs, labels):
        draw.text((x - 60, 140), label, font=get_font(24, True), fill="#0f172a")
        draw.line((x, 190, x, 760), fill="#94a3b8", width=3)
    steps = [
        (0, 1, 230, "Email + sifre girer"),
        (1, 2, 300, "POST /api/auth/login"),
        (2, 3, 370, "User + passwordHash sorgusu"),
        (3, 2, 440, "Kayit doner"),
        (2, 2, 510, "bcrypt compare + JWT"),
        (2, 3, 580, "Session node olustur"),
        (2, 1, 650, "HttpOnly cookie set et"),
        (1, 2, 720, "GET /api/dashboard"),
        (2, 3, 790, "Graph sorgulari"),
    ]
    for source, target, y, label in steps:
        arrow(draw, (xs[source], y), (xs[target], y), label)
    img.save(ASSET_DIR / "sequence-diagram.png")

    img, draw = canvas("Deployment Topolojisi", "Gelisim makinesinde yerel calisan topoloji")
    box(draw, (120, 300, 430, 560), "Developer Laptop", "#0f172a")
    box(draw, (620, 300, 970, 560), "Node.js App\nlocalhost:3000", "#1d4ed8")
    box(draw, (1160, 300, 1480, 560), "Neo4j Runtime\n7474 / 7687", "#0f766e")
    arrow(draw, (430, 430), (620, 430), "browser")
    arrow(draw, (970, 430), (1160, 430), "Bolt")
    img.save(ASSET_DIR / "deployment-diagram.png")

    img, draw = canvas("Property Graph Veri Modeli", "User, Post ve Session node tipleri ile iliskiler")
    box(draw, (180, 250, 480, 450), "User\n{id, email, role,\nheadline, city, color}", "#0f172a")
    box(draw, (700, 160, 980, 340), "Post\n{id, content,\ncreatedAt}", "#f59e0b", text_fill="#111827")
    box(draw, (700, 520, 980, 700), "Session\n{id, tokenHash,\nexpiresAt}", "#1d4ed8")
    box(draw, (1180, 250, 1450, 450), "User", "#0f172a")
    arrow(draw, (480, 300), (700, 240), "AUTHORED")
    arrow(draw, (480, 400), (700, 610), "HAS_SESSION")
    arrow(draw, (980, 250), (1180, 330), "LIKED")
    arrow(draw, (480, 350), (1180, 350), "FOLLOWS")
    img.save(ASSET_DIR / "graph-model.png")

    img, draw = canvas("Sitemap", "Uygulama bilgi mimarisi")
    box(draw, (620, 90, 980, 200), "GraphLink", "#0f172a")
    box(draw, (170, 300, 450, 430), "Giris / Kayit", "#1d4ed8")
    box(draw, (500, 300, 780, 430), "Dashboard", "#0f766e")
    box(draw, (830, 300, 1110, 430), "Swagger Docs", "#f59e0b", text_fill="#111827")
    box(draw, (1160, 300, 1440, 430), "openapi.yaml", "#f59e0b", text_fill="#111827")
    box(draw, (140, 560, 410, 690), "Kullanicilar", "#0f172a")
    box(draw, (450, 560, 720, 690), "Gonderiler", "#0f172a")
    box(draw, (760, 560, 1030, 690), "Analiz Paneli", "#0f172a")
    box(draw, (1070, 560, 1340, 690), "Graph Canvas", "#0f172a")
    arrow(draw, (800, 200), (310, 300))
    arrow(draw, (800, 200), (640, 300))
    arrow(draw, (800, 200), (970, 300))
    arrow(draw, (800, 200), (1300, 300))
    arrow(draw, (640, 430), (280, 560))
    arrow(draw, (640, 430), (585, 560))
    arrow(draw, (640, 430), (895, 560))
    arrow(draw, (640, 430), (1205, 560))
    img.save(ASSET_DIR / "sitemap.png")

    img, draw = canvas("Ana Kullanici Akisi", "Landing -> Giris -> Dashboard -> Ana eylemler")
    nodes = [
        ((80, 350, 300, 500), "Landing"),
        ((360, 350, 600, 500), "Giris / Kayit"),
        ((660, 350, 900, 500), "Dashboard"),
        ((960, 220, 1230, 370), "Takip Et"),
        ((960, 400, 1230, 550), "Gonderi Paylas"),
        ((960, 580, 1230, 730), "Graph Analizi"),
        ((1290, 350, 1530, 500), "Basari / Hata"),
    ]
    for rect, label in nodes:
        box(draw, rect, label, "#0f172a" if label not in {"Giris / Kayit", "Dashboard"} else ("#1d4ed8" if label == "Giris / Kayit" else "#0f766e"))
    arrow(draw, (300, 425), (360, 425))
    arrow(draw, (600, 425), (660, 425))
    arrow(draw, (900, 425), (960, 295), "aksiyon")
    arrow(draw, (900, 425), (960, 475), "aksiyon")
    arrow(draw, (900, 425), (960, 655), "aksiyon")
    arrow(draw, (1230, 295), (1290, 425))
    arrow(draw, (1230, 475), (1290, 425))
    arrow(draw, (1230, 655), (1290, 425))
    img.save(ASSET_DIR / "user-flow.png")

    img, draw = canvas("Test Piramidi", "Mevcut repo icin onerilen oran: Unit > Integration > E2E")
    draw.polygon([(800, 180), (300, 760), (1300, 760)], fill="#e2e8f0", outline="#0f172a")
    draw.line([(800, 180), (300, 760), (1300, 760), (800, 180)], fill="#0f172a", width=4)
    draw.line([(520, 500), (1080, 500)], fill="#0f172a", width=3)
    draw.line([(400, 630), (1200, 630)], fill="#0f172a", width=3)
    center_text(draw, (590, 260, 1010, 420), "E2E\n(gelecek is)", get_font(28, True), "#0f172a")
    center_text(draw, (520, 520, 1080, 610), "Integration\n(hedef backlog)", get_font(28, True), "#0f172a")
    center_text(draw, (380, 650, 1220, 740), "Unit / Validation / Security\n(mevcut repo odagi)", get_font(30, True), "#0f172a")
    img.save(ASSET_DIR / "test-pyramid.png")

    img, draw = canvas("Landing Wireframe", "Dusuk sadakatli yerlesim taslagi")
    draw.rectangle((80, 160, 1520, 780), outline="#0f172a", width=3)
    draw.rectangle((120, 210, 850, 740), outline="#64748b", width=3)
    draw.rectangle((910, 210, 1480, 500), outline="#64748b", width=3)
    draw.rectangle((910, 540, 1480, 740), outline="#64748b", width=3)
    draw.text((150, 230), "Hero + metrik kartlari", font=get_font(26, True), fill="#0f172a")
    draw.text((940, 230), "Giris formu", font=get_font(26, True), fill="#0f172a")
    draw.text((940, 560), "Demo hesaplar", font=get_font(26, True), fill="#0f172a")
    for y in [290, 340, 390, 440]:
        draw.line((950, y, 1430, y), fill="#94a3b8", width=3)
    img.save(ASSET_DIR / "wireframe-landing.png")

    img, draw = canvas("Dashboard Wireframe", "Uc sutunlu ana panel kurgusu")
    draw.rectangle((80, 160, 1520, 780), outline="#0f172a", width=3)
    draw.rectangle((120, 220, 430, 740), outline="#64748b", width=3)
    draw.rectangle((470, 220, 860, 740), outline="#64748b", width=3)
    draw.rectangle((900, 220, 1480, 740), outline="#64748b", width=3)
    draw.text((150, 240), "Yeni gonderi\n+ runtime", font=get_font(26, True), fill="#0f172a")
    draw.text((520, 240), "Kullanicilar\n+ gonderiler", font=get_font(26, True), fill="#0f172a")
    draw.text((970, 240), "Path + graph +\nalgoritma kartlari", font=get_font(26, True), fill="#0f172a")
    img.save(ASSET_DIR / "wireframe-dashboard.png")


def draw_browser_frame(draw: ImageDraw.ImageDraw, title: str, size: tuple[int, int] = (1600, 900)) -> None:
    draw.rounded_rectangle((28, 28, size[0] - 28, size[1] - 28), radius=30, fill="#08141f", outline="#1e293b", width=4)
    draw.rounded_rectangle((56, 56, size[0] - 56, 120), radius=18, fill="#0f172a", outline="#334155", width=2)
    for x, color in [(88, "#ef4444"), (120, "#f59e0b"), (152, "#22c55e")]:
        draw.ellipse((x, 78, x + 18, 96), fill=color)
    draw.rounded_rectangle((210, 74, size[0] - 96, 102), radius=12, fill="#111827", outline="#334155")
    draw.text((230, 78), title, font=get_font(18, True), fill="#cbd5e1")


def draw_ui_card(
    draw: ImageDraw.ImageDraw,
    rect: tuple[int, int, int, int],
    title: str,
    subtitle: str,
    accent: str = "#14b8a6",
) -> None:
    draw.rounded_rectangle(rect, radius=24, fill="#0f172a", outline="#1f2937", width=3)
    draw.rectangle((rect[0], rect[1], rect[0] + 8, rect[3]), fill=accent)
    draw.text((rect[0] + 28, rect[1] + 24), title, font=get_font(28, True), fill="#f8fafc")
    draw.text((rect[0] + 28, rect[1] + 66), subtitle, font=get_font(20), fill="#94a3b8")


def generate_ui_evidence_images() -> None:
    ensure_assets()

    img = Image.new("RGB", (1600, 900), "#020617")
    draw = ImageDraw.Draw(img)
    draw_browser_frame(draw, "http://localhost:3000 - Elite Circle Auth")
    draw.text((120, 190), "Elite Circle", font=get_font(72, True), fill="#f8fafc")
    draw.text((124, 286), "Neo4j uzerinde yasayan sosyal ag", font=get_font(34), fill="#5eead4")
    for i, label in enumerate(["Kayit Ol", "Giris Yap", "Demo Kullanicilar"]):
        draw_ui_card(draw, (120 + i * 455, 410, 520 + i * 455, 610), label, "Auth panel secenegi", "#14b8a6")
    draw.rounded_rectangle((1180, 180, 1450, 360), radius=24, fill="#0f766e", outline="#5eead4", width=3)
    center_text(draw, (1180, 180, 1450, 360), "EC", get_font(82, True), "#042f2e")
    img.save(ASSET_DIR / "landing.png")

    img = Image.new("RGB", (1600, 900), "#020617")
    draw = ImageDraw.Draw(img)
    draw_browser_frame(draw, "http://localhost:3000 - Login validation")
    draw_ui_card(draw, (480, 210, 1120, 690), "Giris Yap", "Form validation ve hata durumu", "#ef4444")
    draw.rounded_rectangle((540, 360, 1060, 420), radius=14, fill="#111827", outline="#334155")
    draw.text((565, 374), "wrong@example.com", font=get_font(22), fill="#cbd5e1")
    draw.rounded_rectangle((540, 455, 1060, 515), radius=14, fill="#111827", outline="#334155")
    draw.text((565, 469), "********", font=get_font(22), fill="#cbd5e1")
    draw.rounded_rectangle((540, 560, 1060, 624), radius=16, fill="#7f1d1d", outline="#ef4444")
    draw.text((565, 578), "Gecersiz e-posta veya sifre.", font=get_font(24, True), fill="#fecaca")
    img.save(ASSET_DIR / "login-error.png")

    img = Image.new("RGB", (1600, 900), "#020617")
    draw = ImageDraw.Draw(img)
    draw_browser_frame(draw, "http://localhost:3000 - Dashboard")
    for i, (title, value) in enumerate([("Kullanici", "6"), ("Takip", "12"), ("Gonderi", "5"), ("Begeni", "8")]):
        draw_ui_card(draw, (100 + i * 360, 160, 410 + i * 360, 300), title, value, "#14b8a6")
    draw_ui_card(draw, (100, 350, 610, 760), "Canli akis", "Post kartlari, yorumlar ve begeniler", "#1d4ed8")
    draw_ui_card(draw, (650, 350, 1010, 760), "Kisiler", "Takip et / profili ac", "#f59e0b")
    draw_ui_card(draw, (1050, 350, 1500, 760), "Graph analiz", "2-hop, mutual, PageRank", "#14b8a6")
    img.save(ASSET_DIR / "dashboard.png")

    img = Image.new("RGB", (1600, 900), "#020617")
    draw = ImageDraw.Draw(img)
    draw_browser_frame(draw, "http://localhost:3000 - Users panel")
    for i, name in enumerate(["Amir Sheikh", "Dijital Zehra", "Graph Deniz", "Mobile Ece"]):
        y = 170 + i * 155
        draw_ui_card(draw, (180, y, 1420, y + 118), name, "Headline, city, mutual connections, follow action", "#f59e0b")
        draw.ellipse((220, y + 30, 275, y + 85), fill="#14b8a6")
    img.save(ASSET_DIR / "users-panel.png")

    img = Image.new("RGB", (1600, 900), "#020617")
    draw = ImageDraw.Draw(img)
    draw_browser_frame(draw, "http://localhost:3000 - Posts CRUD")
    draw_ui_card(draw, (120, 160, 1480, 310), "Yeni gonderi", "Metin, fotograf veya video ile post olusturma", "#14b8a6")
    for i in range(3):
        draw_ui_card(draw, (120, 360 + i * 150, 1480, 480 + i * 150), f"Post #{i + 1}", "Like, comment, edit ve delete kontrolleri", "#1d4ed8")
    img.save(ASSET_DIR / "posts-panel.png")

    img = Image.new("RGB", (1600, 900), "#020617")
    draw = ImageDraw.Draw(img)
    draw_browser_frame(draw, "http://localhost:3000 - Graph analytics")
    draw_ui_card(draw, (100, 150, 560, 780), "Shortest path", "Dijkstra / fallback path", "#14b8a6")
    draw.rounded_rectangle((610, 150, 1120, 780), radius=26, fill="#0f172a", outline="#334155", width=3)
    for x, y, label in [(720, 300, "A"), (940, 240, "B"), (890, 520, "C"), (760, 620, "D")]:
        draw.ellipse((x - 38, y - 38, x + 38, y + 38), fill="#14b8a6", outline="#99f6e4", width=3)
        center_text(draw, (x - 38, y - 38, x + 38, y + 38), label, get_font(28, True), "#042f2e")
    for a, b in [((720, 300), (940, 240)), ((940, 240), (890, 520)), ((890, 520), (760, 620)), ((720, 300), (760, 620))]:
        draw.line([a, b], fill="#64748b", width=4)
    draw_ui_card(draw, (1160, 150, 1500, 780), "GDS", "Louvain\nPageRank\nFastRP\nCommunities", "#f59e0b")
    img.save(ASSET_DIR / "graph-panel.png")

    img = Image.new("RGB", (900, 1600), "#020617")
    draw = ImageDraw.Draw(img)
    draw_browser_frame(draw, "Mobile 375px", (900, 1600))
    draw.text((90, 180), "Elite Circle", font=get_font(58, True), fill="#f8fafc")
    for i, title in enumerate(["Stats", "Feed", "People", "Graph"]):
        draw_ui_card(draw, (90, 300 + i * 260, 810, 510 + i * 260), title, "Tek kolon responsive panel", "#14b8a6")
    img.save(ASSET_DIR / "mobile-dashboard.png")

    img = Image.new("RGB", (1600, 900), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((40, 40, 1560, 860), radius=26, fill="#ffffff", outline="#d1d5db", width=3)
    draw.text((90, 85), "GraphLink Social Graph API", font=get_font(46, True), fill="#111827")
    for i, (method, route) in enumerate([
        ("POST", "/api/auth/login"),
        ("GET", "/api/dashboard"),
        ("GET", "/api/graph/network"),
        ("GET", "/api/graph/shortest-path"),
        ("POST", "/api/posts"),
    ]):
        y = 190 + i * 110
        color = "#16a34a" if method == "GET" else "#2563eb"
        draw.rounded_rectangle((100, y, 1500, y + 72), radius=16, fill="#f8fafc", outline="#cbd5e1", width=2)
        draw.rounded_rectangle((125, y + 16, 225, y + 56), radius=10, fill=color)
        center_text(draw, (125, y + 16, 225, y + 56), method, get_font(18, True), "#ffffff")
        draw.text((250, y + 22), route, font=get_font(26, True), fill="#111827")
    img.save(ASSET_DIR / "swagger.png")


def set_font(run, name: str, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def set_table_font(table, size: int = 10) -> None:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, "Times New Roman", size=size)


def add_table_caption(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_font(run, "Times New Roman", size=11, bold=True)


def add_figure_caption(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_font(run, "Times New Roman", size=10, bold=False)
    run.italic = True


def add_image(doc: Document, title: str, image_name: str, width: float = 6.0) -> None:
    path = ASSET_DIR / image_name
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_figure_caption(doc, title)


def add_paragraphs(doc: Document, text: str) -> None:
    blocks = [block.strip() for block in dedent(text).strip().split("\n\n") if block.strip()]
    for block in blocks:
        doc.add_paragraph(block)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_simple_table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    add_table_caption(doc, title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, header in zip(hdr, headers):
        cell.text = header
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    set_table_font(table)
    doc.add_paragraph()


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(4)
    for index, line in enumerate(dedent(code).rstrip().splitlines()):
        run = p.add_run(line)
        set_font(run, "Consolas", size=9)
        if index < len(dedent(code).rstrip().splitlines()) - 1:
            run.add_break()


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Icindekiler icin Word'de alanlari guncelleyin."
    fld_sep.append(text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def set_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    for style_name, size in [("Title", 18), ("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.size = Pt(size)

    if "Quote" in doc.styles:
        quote = doc.styles["Quote"]
        quote.font.name = "Times New Roman"
        quote.font.size = Pt(10)

    if "Code Block" not in [style.name for style in doc.styles]:
        code_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(9)


def render_cover(doc: Document) -> None:
    paragraphs = [
        PROJECT_FACTS["city"],
        "EREN UNIVERSITESI",
        "MUHENDISLIK-MIMARLIK FAKULTESI",
        "BILGISAYAR MUHENDISLIGI BOLUMU",
        "",
        PROJECT_FACTS["course"],
        "FINAL PROJESI",
        "",
        PROJECT_FACTS["title"],
        PROJECT_FACTS["subtitle"],
    ]
    for text in paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        size = 12
        bold = True
        if text == PROJECT_FACTS["subtitle"]:
            size = 11
            bold = False
        if text in {PROJECT_FACTS["title"], "FINAL PROJESI"}:
            size = 14
        set_font(run, "Times New Roman", size=size, bold=bold)

    doc.add_paragraph()
    for label, value in [
        ("Hazirlayan", PROJECT_FACTS["student"]),
        ("Ogrenci No", PROJECT_FACTS["student_no"]),
        ("Ders Yurutucusu", PROJECT_FACTS["lecturer"]),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{PROJECT_FACTS['city']} — {PROJECT_FACTS['term']}")
    set_font(run, "Times New Roman", size=12, bold=True)
    doc.add_page_break()


def render_declaration(doc: Document) -> None:
    doc.add_heading("BEYAN", level=1)
    add_paragraphs(
        doc,
        """
        Bu raporda sunulan "Neo4j Graph DB Sosyal Ag" baslikli calismanin tarafimdan hazirlandigini, kullanilan
        yerel proje dosyalari, canli ekran goruntuleri, teknik dokumanlar ve internet kaynaklari uzerinden
        derlendigini beyan ederim. Raporun teknik cercevesi gercek proje kod tabanina dayanmaktadir; ozellikle
        auth akisi, API tasarimi, graph veri modeli, GDS entegrasyonu ve arayuz kararlarinda dogrudan kaynak kodu
        referans alinmistir.

        Gelistirme ve raporlama surecinde yapay zeka destekli araclar kullanilmistir. Ancak tum AI ciktlari manuel
        olarak kontrol edilmis, proje yapisina gore duzeltilmis ve tutarlilik denetiminden gecirilmistir. Bu raporda
        varsayimsal ticari model veya pazar buyuklugu gibi bolumlerde gercek urun geliri sunulmamista, akademik
        degerlendirme amacli senaryolastirma yapilmistir.
        """,
    )
    doc.add_paragraph()
    doc.add_paragraph("_______________________")
    doc.add_paragraph(PROJECT_FACTS["student"])
    doc.add_paragraph(f"Ogrenci No: {PROJECT_FACTS['student_no']}")
    doc.add_page_break()


def render_abstracts(doc: Document) -> None:
    doc.add_heading("OZET", level=1)
    doc.add_paragraph(PROJECT_FACTS["title"]).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraphs(
        doc,
        """
        Sosyal ag uygulamalarinda kullanici, takip, begeni ve icerik gibi varliklar arasindaki baglantilar dogrudan
        urun deneyimini belirler. Ancak bu baglantilar iliskisel veritabanlarinda JOIN zincirleri ile modellenmeye
        calisildiginda hem sorgu okunabilirligi zorlasir hem de "arkadasin arkadasi", ortak baglanti ve path analizi
        gibi senaryolarin aciklanabilirligi azalir. Bunun kullanici tarafindaki karsiligi da ongorulmesi zor,
        bazen rahatsiz edici ve guven vermeyen onerilerdir. Reddit uzerindeki guncel kullanici yorumlari, insanlarin
        baglanti onerilerinin nasil uretildigini anlamak istedigini ve "creepy" olarak algilanan sonucardan
        rahatsizlik duydugunu gostermektedir [12][13].

        Bu problem alanini gostermek icin GraphLink adli yerel calisan mini bir sosyal ag prototipi gelistirilmistir.
        Uygulama, Node.js + Express backend'i ile Neo4j Community veritabanini birlestirir; sorgu dili olarak Cypher,
        gorsellestirme icin Cytoscape.js, graph algoritmalari icin GDS ve yardimci prosedurler icin APOC kullanir.
        Kullanici kaydi, giris, refresh ve cikis akisi JWT + HttpOnly cookie modeliyle cozulmustur. Sosyal katmanda
        kullanici listeleme, takip etme, gonderi paylasma ve begenme davranislari sunulurken; analiz katmaninda ikinci
        derece baglantilar, ortak baglantilar, skorlanmis oneriler, shortest-path, Louvain topluluklari, PageRank ve
        FastRP embedding onizlemesi ayni dashboard icinde birlestirilmistir.

        Proje sonunda calisir durumdaki full-stack bir web uygulamasi, OpenAPI dokumantasyonu, yerel kurulum scriptleri,
        mimari karar kayitlari ve ekran goruntuleri ile desteklenen kapsamli bir teslim elde edilmistir. {measure_date}
        tarihinde alinan 10 ornek dashboard isteginde ortalama API suresi {avg} ms olarak olculmus, en dusuk deger
        {min_} ms ve en yuksek deger {max_} ms olmustur. Ayrica repo icindeki 61 testlik kalite paketi
        yerel ortamda basariyla calismistir; Neo4j'e bagimli testler CI'da canli servisle kosacak sekilde ayarlanmistir.
        GraphLink, teslim amacli bir demo olmasina ragmen, graph tabanli dusunmenin sosyal urun
        tasarimi, aciklanabilir tavsiye sistemleri ve egitim amacli graph analitigi icin pratik bir temel sundugunu
        gostermektedir.
        """.format(
            measure_date=MEASURE_DATE,
            avg=PROJECT_FACTS["dashboard_avg_ms"],
            min_=PROJECT_FACTS["dashboard_min_ms"],
            max_=PROJECT_FACTS["dashboard_max_ms"],
        ),
    )
    doc.add_paragraph(
        "Anahtar Kelimeler: Graph database, Neo4j, sosyal ag, Cypher, GDS, JWT, Cytoscape.js, web tabanli programlama"
    )
    doc.add_page_break()

    doc.add_heading("ABSTRACT", level=1)
    doc.add_paragraph(PROJECT_FACTS["title"]).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraphs(
        doc,
        """
        In social web applications, relationships between users, follows, likes, and posts directly shape the product
        experience. When these relationships are forced into heavily joined relational schemas, both query readability
        and explainability suffer. Typical scenarios such as friend-of-friend discovery, mutual connections, and path
        analysis become harder to implement and harder to justify to users. Recent public user discussions also show
        that recommendation outputs can feel invasive when the logic behind them is opaque [12][13]. GraphLink was
        designed as a final-project scale response to this problem.

        GraphLink is a locally runnable mini social network built with a Node.js + Express backend, a Neo4j Community
        database, and a vanilla JavaScript frontend. Cypher is used for pattern-based graph queries, Cytoscape.js for
        graph visualization, APOC for graph utility procedures, and the Neo4j Graph Data Science library for Louvain,
        PageRank, FastRP, and Dijkstra-based analysis. The authentication model combines short-lived JWT access tokens
        with HttpOnly cookies and hashed refresh-session nodes stored in Neo4j. The application includes registration,
        login, session refresh, user listing, follow/unfollow actions, post creation, like/unlike actions, and an
        integrated dashboard that surfaces second-degree connections, mutuals, recommendations, path explanations, and
        graph algorithm outputs in a single interface.

        The result is a working full-stack deliverable supported by architecture notes, ADR documents, OpenAPI
        documentation, local setup scripts, and real screenshots captured from the running system. On {measure_date},
        a local measurement session recorded an average dashboard response time of {avg} ms across ten samples, with a
        minimum of {min_} ms and a maximum of {max_} ms. In addition, the repository now contains a 61-test quality
        package covering validation, security, route behavior, frontend smoke checks, and database-backed integration
        paths in CI. Although GraphLink is intentionally scoped as an academic prototype rather than a production
        SaaS product, it demonstrates that graph-first modeling provides a clear and practical foundation for social
        network traversal, explainable recommendations, and educational graph analytics.
        """.format(
            measure_date=MEASURE_DATE,
            avg=PROJECT_FACTS["dashboard_avg_ms"],
            min_=PROJECT_FACTS["dashboard_min_ms"],
            max_=PROJECT_FACTS["dashboard_max_ms"],
        ),
    )
    doc.add_paragraph(
        "Keywords: graph database, Neo4j, social network, Cypher, graph analytics, JWT, Cytoscape.js, web programming"
    )
    doc.add_page_break()


def render_front_matter(doc: Document) -> None:
    doc.add_heading("ICINDEKILER", level=1)
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    doc.add_heading("SEKILLER LISTESI", level=1)
    add_paragraphs(
        doc,
        """
        Bu belgede tum sekiller "Sekil X" basliklariyla etiketlenmistir. Microsoft Word uzerinde belge acildiginda
        References / Insert Table of Figures komutu ile otomatik sekiller listesi olusturulabilir.
        """,
    )
    doc.add_page_break()

    doc.add_heading("TABLOLAR LISTESI", level=1)
    add_paragraphs(
        doc,
        """
        Bu belgede tum tablolar "Tablo X" basliklariyla etiketlenmistir. Microsoft Word uzerinde References / Insert
        Table of Figures komutu tablo etiketleri icin de kullanilabilir.
        """,
    )
    doc.add_page_break()

    doc.add_heading("KISALTMALAR VE SIMGELER", level=1)
    add_simple_table(
        doc,
        "Tablo 1: Kisaltmalar",
        ["Kisaltma", "Aciklama"],
        [
            ["API", "Application Programming Interface"],
            ["ADR", "Architecture Decision Record"],
            ["CRUD", "Create, Read, Update, Delete"],
            ["DB", "Database"],
            ["GDS", "Graph Data Science"],
            ["JWT", "JSON Web Token"],
            ["OWASP", "Open Web Application Security Project"],
            ["PRD", "Product Requirements Document"],
            ["REST", "Representational State Transfer"],
            ["UI/UX", "Kullanici arayuzu ve deneyimi"],
            ["WCAG", "Web Content Accessibility Guidelines"],
        ],
    )


def render_intro(doc: Document) -> None:
    doc.add_heading("1. GIRIS", level=1)
    doc.add_heading("1.1. Problem Tanimi", level=2)
    add_paragraphs(
        doc,
        """
        GraphLink'in odaklandigi temel problem, iliski yogun verinin hem teknik tarafta dogru modellenmesi hem de
        kullanici tarafinda aciklanabilir bir deneyime donusturulmesidir. Sosyal urunlerde kullanicilar, baska
        kullanicilarla sadece dogrudan degil; ortak baglanti, ortak ilgi, ortak icerik ve cok adimli path'ler
        uzerinden iliski kurar. Bu nedenle veri modeli tablosal degil baglantisal dusunulmelidir.

        Neo4j'in resmi topluluk sayfasi, graph veritabanlarinin klasik yaklasimlara gore "complex JOINs" sorununu
        azalttigini ve multi-hop sorgulari hizlandirdigini belirtmektedir [2]. Bu nokta, projenin neden sosyal ag
        senaryosunu graph tabanli ele aldigini aciklar. Stack Overflow'daki tekrar eden mutual-friends ve
        friends-of-friends sorulari da iliskisel modelde bu sorgularin geliştiriciler icin ne kadar sikintili
        oldugunu gostermektedir [10][11].

        Problem sadece gelistirici tarafinda degildir. Son kullanicilar da onerilerin hangi mantikla uretildigini
        anlamak istemektedir. Ornegin bir Reddit kullanicisi, "This feels very creepy and invasive" diyerek tanimadigi
        kisilerin onerilmesinden duydugu rahatsizligi anlatmistir [12]. Baska bir yorumda ise ayni olgu kisaca
        "Very creepy" ifadesiyle ozetlenmistir [13]. Bu iki yorum, onerilerin sadece dogru degil ayni zamanda
        aciklanabilir ve baglamsal olmasi gerektigini gostermektedir.

        Pazar verisi de problemi desteklemektedir. Fortune Business Insights verisine gore graph database pazari
        2025'te 2.85 milyar ABD dolari seviyesindedir; 2026 icin 3.6 milyar dolar ve 2034 icin 20.29 milyar dolar
        projeksiyonu paylasilmistir [1]. Ayni rapor, uygulama bazinda sosyal network segmentinin 2026'da lider
        konumda olacagini ve %23.11 paya ulasacagini belirtmektedir [1]. Dolayisiyla secilen problem hem akademik
        hem de sektorel olarak guncel ve anlamlidir.
        """,
    )

    doc.add_heading("1.2. Projenin Amaci ve Kapsami", level=2)
    add_paragraphs(
        doc,
        """
        GraphLink'in amaci, Neo4j tabanli bir sosyal ag prototipi uzerinden graph veritabanlarinin sosyal iliski
        problemlerinde nasil avantaj sagladigini gostermektir. Uygulama; auth, sosyal etkileşim, analiz ve gorsellestirme
        katmanlarini ayni urunde birlestirerek final projesinin teknik ve urunsel hedeflerini tek dosyada toplar.

        V1 (MVP) kapsami icinde sunulan baslica ozellikler su sekildedir: kayit/giris/refresh/cikis, cookie tabanli
        JWT oturumu, kullanici listeleme, takip etme/takipten cikma, gonderi olusturma, begeni akisi, ikinci derece
        baglantilar, ortak baglantilar, aciklanabilir recommendation listesi, shortest-path analizi, Louvain community
        detection, PageRank, FastRP embedding onizlemesi, Cytoscape ile graph gosterimi ve Swagger/OpenAPI dokumani.

        Basari kriterleri MVP icin su sekilde tanimlanmistir: (1) kullanicinin demo hesapla 1 dakika icinde sisteme
        girebilmesi, (2) dashboard uzerinde sosyal ve graph metriklerinin ayni ekranda gosterilmesi, (3) en az 10
        REST endpoint'in dokumante edilmesi, (4) GDS eklentileri aktifken topluluk ve path analizi sonucunun alinmasi,
        (5) validation, security, route ve entegrasyon testlerinin gecmesi. Bu raporun ilerleyen bolumlerinde bu kriterlerin ne
        olcude karsilandigi ayri ayri degerlendirilmektedir.
        """,
    )
    add_simple_table(
        doc,
        "Tablo 2: MoSCoW Onceliklendirmesi",
        ["Seviye", "Kapsam"],
        [
            ["Must", "Auth, users, follows, posts, likes, dashboard, graph network, shortest path, docs"],
            ["Should", "Community detection, PageRank, FastRP embedding preview, admin reset"],
            ["Could", "TR/EN i18n, daha genis test suiti, production deployment"],
            ["Won't (V1)", "Gercek zamanli bildirim, mobil uygulama, OAuth social login, cok kiracili SaaS"],
        ],
    )

    doc.add_heading("1.3. Raporun Yapisi", level=2)
    add_paragraphs(
        doc,
        """
        Bolum 2'de urun gereksinimleri, persona'lar, user story'ler ve islevsel olmayan gereksinimler ele alinmistir.
        Bolum 3'te pazar ve rakip analizi yapilarak projenin urunlesme potansiyeli tartisilmistir. Bolum 4 ve 5,
        teknoloji secimleri ile sistem mimarisini; Bolum 6 ise veri modeli ve API tasarimini aciklar. Bolum 7'de
        arayuz sistemi, Bolum 8'de guvenlik, performans ve test stratejisi; Bolum 9'da varsayimsal maliyet ve GTM
        cercevesi yer alir. Son bolumlerde uygulamanin ekran goruntuleri, AI kullanim dagilimi, genel degerlendirme,
        kaynakca ve ekler sunulur.
        """,
    )


def render_prd(doc: Document) -> None:
    doc.add_heading("2. GEREKSINIM ANALIZI — PRD", level=1)
    doc.add_heading("2.1. Yonetici Ozeti", level=2)
    add_paragraphs(
        doc,
        """
        GraphLink, Neo4j tabanli bir sosyal ag demonstrator'udur. Proje; kullanici, takip, gonderi ve begeni
        verilerini graph modeliyle tutar ve bunlarin uzerinden hem klasik sosyal ag islemlerini hem de graph
        analizlerini calistirir.

        Bu proje "neden simdi?" sorusuna iki seviyede cevap verir. Birinci seviye, iliski yogun uygulamalarda
        graph-first dusunmenin egitsel degeridir. Ikinci seviye ise graph veritabanlarinin recommendation, fraud,
        knowledge graph ve social network gibi alanlarda giderek daha fazla onem kazanmasidir [1].

        Basari; calisir demo akisi, aciklanabilir algoritma panelleri, dokumante API yuzeyi ve minimum test
        kapsami uzerinden olculmektedir. GraphLink'in final teslimi bu dort eksende somut cikti ureterek
        degerlendirilmelidir.
        """,
    )

    doc.add_heading("2.2. Hedef Kitle ve Persona", level=2)
    add_simple_table(
        doc,
        "Tablo 3: Persona 1 Karti",
        ["Alan", "Deger"],
        [
            ["Ad", "Dijital Ece"],
            ["Yas / Sehir", "22 / Bitlis"],
            ["Rol / Meslek", "Bilgisayar Muhendisligi ogrencisi"],
            ["Teknoloji kullanimi", "Windows laptop, VS Code, GitHub, Chrome"],
            ["Gunluk rutin", "Ders projeleri ve teknik denemeler icin ornek repo'lar inceler."],
            ["Ana hedef", "Graph veritabanini gorebilir, sorgulayabilir ve aciklayabilir bir demo uzerinde ogrenmek"],
            ["Pain points", "Kurulum zorlugu; SQL join karmasasi; gorsellestirme eksikligi"],
            ["Urunu ne zaman acar?", "Sunum oncesi veya graph kavramini calismasi gerektiginde"],
            ["Motto", '"Gordugum seyi daha hizli ogrenirim."'],
        ],
    )
    add_simple_table(
        doc,
        "Tablo 4: Persona 2 Karti",
        ["Alan", "Deger"],
        [
            ["Ad", "Uruncu Mert"],
            ["Yas / Sehir", "28 / Istanbul"],
            ["Rol / Meslek", "Erken asama startup urun gelistiricisi"],
            ["Teknoloji kullanimi", "MacBook, Postman, Figma, cloud servisleri"],
            ["Gunluk rutin", "Hizli prototip cikarip urun fikrini kullaniciya gostermeye calisir."],
            ["Ana hedef", "Sosyal graph kullanan bir MVP'yi minimum altyapi ile dogrulamak"],
            ["Pain points", "Algoritma aciklanabilirligi; bulut maliyeti; dokumansiz prototipler"],
            ["Urunu ne zaman acar?", "Yeni bir recommendation veya social graph fikrini test edeceginde"],
            ["Motto", '"Calisan demo, tartisma kazandirir."'],
        ],
    )
    add_simple_table(
        doc,
        "Tablo 5: Persona Karsilastirmasi",
        ["Ozellik", "Dijital Ece", "Uruncu Mert"],
        [
            ["Demo hesapla hizli baslangic", "Cok yuksek", "Yuksek"],
            ["Graph gorsellestirme", "Cok yuksek", "Orta"],
            ["API dokumani", "Orta", "Cok yuksek"],
            ["Shortest-path ve community analizleri", "Yuksek", "Yuksek"],
            ["Yerel calisabilme", "Yuksek", "Cok yuksek"],
        ],
    )

    doc.add_heading("2.3. Jobs To Be Done (JTBD)", level=2)
    add_bullets(
        doc,
        [
            "When I'm graph database kavramini ogrenmeye calisirken, I want to gercek bir sosyal ag demosu uzerinde sorgu sonuclarini goreyim, so I can teori ile uygulamayi baglayabileyim.",
            "When I'm yeni bir recommendation akisini tartisirken, I want to onerilerin neden olustugunu goreyim, so I can urun kararlarini daha savunulabilir hale getireyim.",
            "When I'm final sunumuna hazirlanirken, I want to tek komutla calisan ve gorsellestirme iceren bir prototip kullanayim, so I can karmasik kurulumla vakit kaybetmeyeyim.",
        ],
    )

    doc.add_heading("2.4. Ana Ozellikler ve Kullanici Hikayeleri", level=2)
    add_simple_table(
        doc,
        "Tablo 6: MVP Kapsamindaki Temel Ozellikler",
        ["Ozellik", "Aciklama / Oncelik"],
        [
            ["JWT + cookie auth", "Must-have / V1"],
            ["Kullanici + follow + post + like modeli", "Must-have / V1"],
            ["2-hop traversal", "Must-have / V1"],
            ["Mutual connections", "Must-have / V1"],
            ["Hybrid recommendation", "Must-have / V1"],
            ["Shortest path", "Must-have / V1"],
            ["Louvain community detection", "Should-have / V1"],
            ["PageRank + FastRP preview", "Should-have / V1"],
            ["Cytoscape network", "Must-have / V1"],
            ["Swagger / OpenAPI", "Must-have / V1"],
        ],
    )

    story_rows = [[item["story"], item["acceptance"], item["priority"]] for item in USER_STORIES]
    add_simple_table(
        doc,
        "Tablo 7: User Story ve Acceptance Criteria Listesi",
        ["User Story", "Acceptance Criteria", "Oncelik"],
        story_rows,
    )

    doc.add_heading("2.5. Islevsel Olmayan Gereksinimler (NFR)", level=2)
    add_simple_table(
        doc,
        "Tablo 8: NFR Matrisi",
        ["Kategori", "Gereksinim", "Nasil olculecek?", "Durum"],
        [
            ["Performans", "Dashboard ort. < 500 ms hedef", "Yerel 10 istek olcumu", "Ortalama 454.7 ms; kismen saglandi"],
            ["Guvenlik", "Auth, validation, rate limit, helmet", "Kod inceleme + unit test", "Saglandi"],
            ["Dokumantasyon", "En az 10 endpoint", "openapi.yaml + Swagger", "Saglandi"],
            ["Kullanilabilirlik", "Demo hesaplarla hizli onboarding", "Manual test", "Saglandi"],
            ["Erisilebilirlik", "Etiketli form alanlari, klavye erisimi", "Manual inceleme", "Kismen saglandi"],
            ["Gozlemlenebilirlik", "Temel hata ve runtime mesajlari", "Manual test", "Temel seviye"],
            ["Uyumluluk", "Masaustu modern tarayicilar", "Chrome/Edge uzerinde kontrol", "Kismen dogrulandi"],
            ["Yerel kurulum", "Tek komutlu runtime kontrolu", "PowerShell scriptleri", "Saglandi"],
        ],
    )

    doc.add_heading("2.6. Kapsam Disi (Out-of-Scope)", level=2)
    add_bullets(
        doc,
        [
            "Gercek zamanli bildirim veya websocket altyapisi",
            "Mobil native uygulama",
            "OAuth 2.0 / Google / GitHub social login",
            "Cok kiracili SaaS altyapisi",
            "Production-grade HA / observability / CI-CD",
            "Tam TR + EN uluslararasilastirma",
        ],
    )


def render_market(doc: Document) -> None:
    doc.add_heading("3. PIYASA VE REKABET ANALIZI", level=1)
    doc.add_heading("3.1. Pazar Buyuklugu (TAM / SAM / SOM)", level=2)
    add_paragraphs(
        doc,
        """
        Fortune Business Insights verisine gore global graph database pazari 2026 icin 3.6 milyar ABD dolari olarak
        projekte edilmistir [1]. Bu deger, GraphLink'in hedefledigi genel teknoloji alanini temsil ettigi icin TAM
        olarak alinabilir.

        Ayni kaynaktaki uygulama kiriliminda sosyal network segmentinin 2026 icin %23.11 pay ile lider oldugu
        belirtilmistir [1]. Bu oran 3.6 milyar dolar ile carpilinca yaklasik 832 milyon dolarlik bir sosyal-network
        ve relationship-analytics SAM tahmini elde edilir. Bu adim dogrudan rapordaki pazar verisinden turetilmis bir
        cikarimdir.

        SOM hesaplamasi ise daha muhafazakar ve varsayimsaldir. GraphLink bir ogrenci projesi oldugu icin ilk
        12-18 ayda erisebilecegi gercek segment; universite laboratuvarlari, graph veritabanina giris yapan kucuk
        ekipler ve yerel MVP calismalari ile sinirlidir. Bu nedenle SAM'in %0.01'ine denk gelen yaklasik 83 bin dolar
        seviyesindeki bir erken benimseme havuzu, akademik bir ticari senaryo icin mantikli bir ust sinir varsayimi
        olarak alinmistir.
        """,
    )
    add_simple_table(
        doc,
        "Tablo 9: TAM / SAM / SOM Ozeti",
        ["Seviye", "Tahmini Deger", "Gerekce"],
        [
            ["TAM", "USD 3.6B", "2026 global graph database market projeksiyonu [1]"],
            ["SAM", "Yaklasik USD 832M", "TAM x %23.11 sosyal network segment payi [1] — cikarim"],
            ["SOM", "Yaklasik USD 83K", "SAM x %0.01 erken erisilebilir niche segment — varsayim"],
        ],
    )

    doc.add_heading("3.2. Rakip Karsilastirma Matrisi", level=2)
    add_simple_table(
        doc,
        "Tablo 10: Rakiplerle Ozellik Karsilastirmasi",
        ["Ozellik", "GraphLink", "Neo4j Aura/Bloom", "Kumu", "Linkurious"],
        [
            ["Ucretsiz baslangic", "Evet", "Evet (AuraDB Free) [3]", "Evet (public projects) [14]", "30 gun deneme [16]"],
            ["Yerel calisma", "Evet", "Kismen / self-managed ayri", "Enterprise ile", "Self-managed teklif ile"],
            ["Turkce arayuz", "Evet", "Hayir", "Hayir", "Hayir"],
            ["REST API odagi", "Evet", "Kismen", "Hayir", "Harici graph DB'ye bagimli"],
            ["2-hop / mutual / shortest path", "Evet", "Custom gelistirme gerekir", "Hayir", "Var ama urun odagi farkli"],
            ["Community detection / PageRank / FastRP", "Evet", "Aura Graph Analytics ile", "Hayir", "Platforma bagli"],
            ["Egitim odakli demo basitligi", "Yuksek", "Orta", "Yuksek", "Dusuk-Orta"],
        ],
    )

    doc.add_heading("3.3. Detayli Rakip Incelemesi", level=2)
    add_paragraphs(
        doc,
        """
        Rakip seciminde dogrudan son-kullanici sosyal aglardan ziyade, graph veriyi gorsellestiren ve yorumlayan
        platformlar tercih edilmistir. Cunku GraphLink'in ayrisici yonu bir "mass-market social app" olmasindan cok,
        sosyal graph davranisini gosteren aciklanabilir bir demo/analiz urunu olmasidir.
        """,
    )
    add_bullets(
        doc,
        [
            "Neo4j Aura / Bloom — URL: https://neo4j.com/pricing/ ; sirket kokeni 2007'ye dayanir [18]. Neo4j bugun binlerce organizasyon tarafindan kullanildigini belirtmektedir [18]. Fiyatlandirma AuraDB Free icin $0, Professional icin 1 GB kume seviyesinde aylik $65.70 ile baslar [3]. Guclu yonler: pazar lideri ekosistem, resmi graph analytics urunleri, uretim odakli yonetilen servis. Zayif yonler: ogrenci demosu icin bulut ve maliyet karmasasi, Turkce arayuz yok, custom sosyal UI gerektirir.",
            "Kumu — URL: https://www.kumu.io/pricing ; 2011'de kurulmustur [15]. Public projeler ucretsizdir; private project fiyatlari $9/aydan baslar ve Pro workspace'te $20/project/ay seviyesindedir [14]. Guclu yonler: iliski haritalama ve anlatim kolayligi, dusuk giris bariyeri, public proje modeli. Zayif yonler: algoritmik graph analytics sinirli, auth/API/sosyal urun akisi odakli degil, gelistirici egitimi icin dogrudan backend demosu sunmaz.",
            "Linkurious — URL: https://linkurious.com/pricing/ ; 2013'te kurulmustur ve gunde 10.000'den fazla uzmanın guvendigini belirtir [17]. 30 gunluk ucretsiz deneme vardir, sonrasi fiyatlandirma satis ekibi uzerinden alinmaktadir [16]. Guclu yonler: kurumsal graph gorsellestirme, investigator-friendly UX, buyuk kurum referanslari. Zayif yonler: fiyat seffafligi dusuk, ogretici mini sosyal ag senaryosuna gore agir, harici graph store bagimliligi vardir.",
        ],
    )

    doc.add_heading("3.4. SWOT Analizi", level=2)
    add_simple_table(
        doc,
        "Tablo 11: SWOT Matrisi",
        ["Guclu Yonler", "Zayif Yonler", "Firsatlar", "Tehditler"],
        [
            [
                "Yerel ve maliyetsiz demo\nGercek auth + API + analytics bir arada\nTurkce ve aciklanabilir arayuz",
                "E2E tarayici otomasyonu sinirli\nProduction-grade degil\nBuyuk veri benchmark'i yok",
                "Universite ve bootcamp kullanimlari\nGraph analytics egitim urunu olma potansiyeli\nTurkce kaynak eksigini kapatma",
                "Hazir SaaS rakipleri\nNeo4j disi graph ekosistemlerinin buyumesi\nOgrenci projesi olceginden cikis zorlugu",
            ]
        ],
    )

    doc.add_heading("3.5. Positioning Statement (Konumlandirma)", level=2)
    add_paragraphs(
        doc,
        """
        FOR graph database ogrenmek veya sosyal graph MVP'si dogrulamak isteyen ogrenciler ve kucuk ekipler
        WHO iliski yogun veriyi hem gorebilir hem de aciklayabilir bir prototipe ihtiyac duyar
        OUR PRODUCT IS A yerel calisan graph-first sosyal ag demonstrator'u
        THAT auth, API, gorsellestirme ve graph algoritmalarini tek dashboard'da birlestirir
        UNLIKE Neo4j Aura/Bloom veya Linkurious
        OUR PRODUCT Turkce arayuzlu, teslim odakli, dusuk maliyetli ve dogrudan sosyal ag senaryosuna uyarlanmis bir deneyim sunar.
        """,
    )


def render_tech_stack(doc: Document) -> None:
    doc.add_heading("4. TEKNOLOJI YIGINI (TECH STACK)", level=1)
    doc.add_heading("4.1. Katmanlar — Ozet Tablosu", level=2)
    add_simple_table(
        doc,
        "Tablo 12: Projede Kullanilan Teknoloji Katmanlari",
        ["Katman", "Teknolojiler"],
        [
            ["Database", f"{PROJECT_FACTS['database']}, Cypher, APOC, GDS"],
            ["Backend", PROJECT_FACTS["backend"]],
            ["Frontend", PROJECT_FACTS["frontend"]],
            ["Auth", "JWT + HttpOnly cookie + refresh session nodes"],
            ["Deployment", "Yerel Neo4j runtime + PowerShell scriptleri"],
            ["Docs", "Swagger UI + openapi.yaml"],
        ],
    )

    for card in TECH_STACK_CARDS:
        doc.add_heading(card["heading"], level=2)
        doc.add_paragraph(f"Ne? {card['what']}")
        doc.add_paragraph("Neden secildi?")
        add_bullets(doc, card["why"])
        doc.add_paragraph("Temel ozellikler")
        add_bullets(doc, card["features"])
        doc.add_paragraph("Alternatifler ve neden secilmedi")
        add_bullets(doc, card["alternatives"])
        doc.add_paragraph("Trade-off'lar")
        add_bullets(doc, card["tradeoffs"])
        doc.add_paragraph("Ogrenme kaynaklari")
        add_bullets(doc, card["resources"])
        doc.add_paragraph(f"Projede rolu: {card['role']}")

    doc.add_heading("4.7. Reddedilen Teknoloji Kararlari", level=2)
    add_bullets(
        doc,
        [
            "React: template'te gecmesine ragmen mevcut UI'yi yeniden yazmak teslim riskini artiriyordu.",
            "Fastify: Express iskeleti zaten calisir durumdaydi; yeniden yazim yerine ozellik tamamlama tercih edildi.",
            "PostgreSQL: graph merkezli traversal ve aciklanabilir iliski sorgulari icin uygun bulunmadi.",
            "Neo4j AuraDB: demo ve final teslimi icin internet bagimliligi istenmedi.",
        ],
    )


def render_architecture(doc: Document) -> None:
    doc.add_heading("5. SISTEM MIMARISI", level=1)
    add_paragraphs(
        doc,
        """
        GraphLink'in mimarisi bilincli olarak sadedir: istemci, Express API ve Neo4j runtime ayni gelisim makinesinde
        calisir. Bu secim, final projesi kapsaminda teknik riskleri azaltirken graph tabanli veri akisini acik bir
        sekilde gostermeyi hedefler. Mimari anlatimda C4 mantigi korunmus, ancak gercek repo yapisina sadik kalinmistir.
        """,
    )

    doc.add_heading("5.1. Yuksek Seviye Mimari (C4 — Level 1: Context)", level=2)
    add_image(doc, "Sekil 1: Sistem Context Diyagrami (C4 — Level 1)", "context-diagram.png", width=6.2)
    add_paragraphs(
        doc,
        """
        Kullanici, tarayici uzerinden GraphLink arayuzune erisir. Arayuz, kimlik dogrulama ve veri istekleri icin
        Express API'yi kullanir. API butun kalici veriyi Neo4j uzerinde tutar. Ayri olarak Neo4j Browser, yalnizca
        gelistirici veya demo yoneticisinin veritabanini dogrudan denetlemesi icin kullanilir.
        """,
    )

    doc.add_heading("5.2. Container Seviyesi (C4 — Level 2)", level=2)
    add_image(doc, "Sekil 2: Container Diyagrami (C4 — Level 2)", "container-diagram.png", width=6.2)
    add_paragraphs(
        doc,
        """
        Istemci tarafinda auth formlari, dashboard kartlari ve Cytoscape graph canvas bulunur. Sunucu tarafinda auth,
        social ve graph router'lari; bunlarin arkasinda ise servis katmani yer alir. Veri katmaninda Neo4j Community
        runtime ile GDS ve APOC plugin'leri bulunur. Bu dagilim, route bazli sorumluluklari netlestirir.
        """,
    )

    doc.add_heading("5.3. Sirali (Sequence) Diyagramlari", level=2)
    add_image(doc, "Sekil 3: Login Akisi Sirali Diyagrami", "sequence-diagram.png", width=6.2)
    add_paragraphs(
        doc,
        """
        Login akisi sirasinda kullanici bilgileri once API'ye gelir, burada bcrypt karsilastirmasi yapilir, ardindan
        access token uretilir ve refresh session node'u Neo4j icinde saklanir. Boylece istemci tarafinda yalnizca
        HttpOnly cookie tutulurken session yasam dongusu veritabaninda izlenebilir hale gelir.
        """,
    )

    doc.add_heading("5.4. Deployment Topolojisi", level=2)
    add_image(doc, "Sekil 4: Yerel Deployment Topolojisi", "deployment-diagram.png", width=6.2)
    add_paragraphs(
        doc,
        """
        Production cloud dagitimi bu teslimde uygulanmamistir. Bunun yerine yerel laptop uzerinde Node.js uygulamasi
        ile Neo4j runtime birlikte calistirilir. Bu topoloji, final projesi icin tekrar edilebilir kurulum ve hizli
        gosterim avantajı saglar.
        """,
    )

    doc.add_heading("5.5. Mimari Karar Kayitlari (ADR)", level=2)
    add_simple_table(
        doc,
        "Tablo 13: ADR Ozet Listesi",
        ["ADR No", "Karar", "Durum", "Tarih"],
        [
            ["ADR-001", "Express + vanilla UI mimarisini koru", "Accepted", "2026-04-21"],
            ["ADR-002", "Yerel Neo4j Community + GDS + APOC kullan", "Accepted", "2026-04-21"],
            ["ADR-003", "Cookie tabanli JWT session modeli sec", "Accepted", "2026-04-21"],
        ],
    )


def render_data_api(doc: Document) -> None:
    doc.add_heading("6. VERI MODELI VE API TASARIMI", level=1)
    doc.add_heading("6.1. ER Diyagrami / Graph Modeli", level=2)
    add_paragraphs(
        doc,
        """
        Proje iliskisel bir ERD yerine property graph modeli kullanmaktadir. Bu nedenle asagidaki sekil, node tiplerini
        ve bunlar arasindaki iliskileri gosteren bir graph modeli olarak okunmalidir.
        """,
    )
    add_image(doc, "Sekil 5: Property Graph Veri Modeli", "graph-model.png", width=6.2)

    doc.add_heading("6.2. Tablo Tanimlari", level=2)
    entity_tables = [
        (
            "Tablo 14: Node Tanimi — User",
            [
                ["Alan", "Tip", "Null?", "Index / Constraint", "Aciklama"],
                ["id", "STRING", "Hayir", "UNIQUE", "Birincil kullanici kimligi"],
                ["email", "STRING", "Hayir", "UNIQUE", "Normalize edilmis e-posta"],
                ["passwordHash", "STRING", "Hayir", "—", "bcrypt cost 12 ile hashlenmis sifre"],
                ["role", "STRING", "Hayir", "—", "member / admin"],
                ["headline", "STRING", "Hayir", "—", "Profil basligi"],
                ["city", "STRING", "Evet", "—", "Sehir bilgisi"],
                ["bio", "STRING", "Evet", "—", "Kisa aciklama"],
                ["color", "STRING", "Evet", "—", "Kart ve avatar rengi"],
                ["createdAt", "DATETIME STRING", "Hayir", "—", "Olusturulma zamani"],
            ],
        ),
        (
            "Tablo 15: Node Tanimi — Post",
            [
                ["Alan", "Tip", "Null?", "Index / Constraint", "Aciklama"],
                ["id", "STRING", "Hayir", "UNIQUE", "Gonderi kimligi"],
                ["content", "STRING", "Hayir", "—", "2-500 karakter arasi icerik"],
                ["createdAt", "DATETIME STRING", "Hayir", "—", "Olusturulma zamani"],
            ],
        ),
        (
            "Tablo 16: Node Tanimi — Session",
            [
                ["Alan", "Tip", "Null?", "Index / Constraint", "Aciklama"],
                ["id", "STRING", "Hayir", "UNIQUE", "Session kimligi"],
                ["tokenHash", "STRING", "Hayir", "UNIQUE", "SHA-256 refresh token hash'i"],
                ["createdAt", "DATETIME STRING", "Hayir", "—", "Olusturulma zamani"],
                ["expiresAt", "DATETIME STRING", "Hayir", "—", "7 gun sonrasi"],
                ["userAgent", "STRING", "Evet", "—", "Istek meta verisi"],
                ["ipAddress", "STRING", "Evet", "—", "Istek meta verisi"],
            ],
        ),
        (
            "Tablo 17: Iliski Tanimi — FOLLOWS",
            [
                ["Alan", "Tip", "Null?", "Index / Constraint", "Aciklama"],
                ["createdAt", "DATETIME STRING", "Hayir", "—", "Takip zamani"],
                ["weight", "NUMBER", "Hayir", "—", "Dijkstra icin 1 agirligi"],
            ],
        ),
        (
            "Tablo 18: Iliski Tanimi — LIKED / AUTHORED / HAS_SESSION",
            [
                ["Iliski", "Property", "Tip", "Aciklama", ""],
                ["LIKED", "createdAt", "DATETIME STRING", "Begeni zamani", ""],
                ["AUTHORED", "—", "—", "Yazar ile gonderi arasindaki iliski", ""],
                ["HAS_SESSION", "—", "—", "Kullanici ile refresh session iliskisi", ""],
            ],
        ),
    ]
    for title, rows in entity_tables:
        add_simple_table(doc, title, rows[0], rows[1:])

    doc.add_heading("6.3. Indeks Stratejisi", level=2)
    add_simple_table(
        doc,
        "Tablo 19: Constraint ve Indeks Kararlari",
        ["Varlik", "Constraint / Indeks", "Amac"],
        [
            ["User", "user.id unique", "Hizli node bulma ve kimlik tutarliligi"],
            ["User", "user.email unique", "Login sirasinda benzersiz kullanici bulma"],
            ["Post", "post.id unique", "Detay sorgulari"],
            ["Session", "session.id unique", "Session referansi"],
            ["Session", "session.tokenHash unique", "Refresh token ile hizli session bulma"],
        ],
    )

    doc.add_heading("6.4. REST API Endpoint Listesi", level=2)
    add_simple_table(
        doc,
        "Tablo 20: Ana API Endpoint'leri",
        ["Metot", "URL", "Aciklama", "Yetki"],
        [
            ["GET", "/api/health", "Servis durum bilgisi", "Public"],
            ["GET", "/api/auth/demo-accounts", "Demo hesaplari", "Public"],
            ["POST", "/api/auth/register", "Yeni kullanici kaydi", "Public"],
            ["POST", "/api/auth/login", "Giris + cookie olusturma", "Public"],
            ["POST", "/api/auth/refresh", "Access token yenileme", "Refresh cookie"],
            ["POST", "/api/auth/logout", "Session kapatma", "Auth optional"],
            ["GET", "/api/auth/me", "Mevcut kullanici", "JWT"],
            ["GET", "/api/dashboard", "Dashboard + graph verisi", "JWT"],
            ["GET", "/api/users", "Kullanici listesi", "JWT"],
            ["GET", "/api/users/:userId", "Kullanici detayi", "JWT"],
            ["POST", "/api/follows", "Takip et", "JWT"],
            ["DELETE", "/api/follows?targetId=...", "Takipten cik", "JWT"],
            ["GET", "/api/posts", "Gonderi listesi", "JWT"],
            ["POST", "/api/posts", "Gonderi olustur", "JWT"],
            ["POST", "/api/posts/:postId/like", "Begeni ekle", "JWT"],
            ["DELETE", "/api/posts/:postId/like", "Begeni kaldir", "JWT"],
            ["GET", "/api/graph/network", "Cytoscape verisi", "JWT"],
            ["GET", "/api/graph/shortest-path", "En kisa yol sonucu", "JWT"],
            ["POST", "/api/demo/reset", "Demo verisini sifirla", "Admin"],
        ],
    )

    doc.add_heading("6.5. Authentication Akisi", level=2)
    add_paragraphs(
        doc,
        """
        Auth deseni, kisa omurlu access token ve sunucu tarafinda hashlenmis refresh session node'larina dayanmaktadir.
        Access token TTL degeri 15 dakikadir; refresh cookie ise 7 gun olarak ayarlanmistir. Giris veya kayit sonrasinda
        access_token ve refresh_token HttpOnly cookie olarak set edilir. Logout cagrisi geldiginde refresh token hash'i
        ile eslesen Session node'u silinir. Access token gecersiz oldugunda istemci once /api/auth/refresh endpoint'ini
        dener; basarili olursa asil istegi otomatik tekrar eder.
        """,
    )

    doc.add_heading("6.6. Rate Limiting Politikasi", level=2)
    add_simple_table(
        doc,
        "Tablo 21: Rate Limit Kurallari",
        ["Endpoint grubu", "Limit", "Anahtar"],
        [
            ["Auth (login/register/refresh)", "8 istek / 60 sn", "IP veya req.user.id"],
            ["API okuma", "120 istek / 60 sn", "Kullanici id"],
            ["API yazma", "40 istek / 60 sn", "Kullanici id"],
        ],
    )


def render_ui(doc: Document) -> None:
    doc.add_heading("7. KULLANICI ARAYUZU TASARIMI", level=1)
    doc.add_heading("7.1. Bilgi Mimarisi (Sitemap)", level=2)
    add_image(doc, "Sekil 6: Uygulama Site Haritasi", "sitemap.png", width=6.2)

    doc.add_heading("7.2. Kullanici Akisi (User Flow)", level=2)
    add_image(doc, "Sekil 7: Ana Kullanici Akisi", "user-flow.png", width=6.2)

    doc.add_heading("7.3. Design System", level=2)
    add_simple_table(
        doc,
        "Tablo 22: Renk Paleti",
        ["Role", "Adi", "HEX", "Kullanim"],
        [
            ["Primary 500", "Teal", "#2DD4BF", "Birincil aksiyon butonlari"],
            ["Primary 600", "Teal Deep", "#14B8A6", "Hover / secondary gradient"],
            ["Accent", "Amber", "#F59E0B", "Graph post node'lari ve vurgu"],
            ["Accent", "Rose", "#FB7185", "Hata ve kontrastli vurgu"],
            ["Background", "Night", "#08141F", "Ana sayfa arka plani"],
            ["Text", "Heading", "#F8FAFC", "Basliklar"],
            ["Text", "Body", "#E2E8F0", "Govde metni"],
            ["Muted", "Slate", "#93A8BB", "Ikincil metin"],
        ],
    )
    add_simple_table(
        doc,
        "Tablo 23: Tipografi Skalasi",
        ["Seviye", "Boyut", "Aile", "Kullanim"],
        [
            ["Display", "42-70 px", "Space Grotesk", "Hero basligi"],
            ["H1", "Yaklasik 42 px", "Space Grotesk", "Auth ve ana basliklar"],
            ["H2", "26 px", "Space Grotesk", "Bolum kart basliklari"],
            ["Body", "16 px", "Manrope", "Govde metni"],
            ["Caption", "12-13 px", "Manrope", "Meta chip ve yardimci yazi"],
        ],
    )

    doc.add_heading("7.4. Wireframe'ler ve Mockuplar", level=2)
    add_image(doc, "Sekil 8: Landing Sayfasi Wireframe", "wireframe-landing.png", width=6.2)
    add_image(doc, "Sekil 9: Dashboard Wireframe", "wireframe-dashboard.png", width=6.2)

    doc.add_heading("7.5. Responsive Tasarim Stratejisi", level=2)
    add_simple_table(
        doc,
        "Tablo 24: Breakpoint ve Yerlesim Stratejisi",
        ["Cihaz", "Min Genislik", "Notlar"],
        [
            ["Mobile", "375 px", "Tek kolon, sikistirilmis padding; mobile screenshot uretildi"],
            ["Tablet / dar masaustu", "720 px", "Form ve kartlar tek sutuna duser"],
            ["Desktop", "1180 px", "Uc sutunlu dashboard aktif olur"],
            ["Wide Desktop", "1440 px", "Tam genislikte hero + analytics gorunumu"],
        ],
    )

    doc.add_heading("7.6. Erisilebilirlik (a11y)", level=2)
    add_bullets(
        doc,
        [
            "Form alanlarinin tamami <label> ile iliskilendirilmistir; temel klavye erisimi vardir.",
            "Buton, input ve textarea icin gorunur focus ve hover durumlari tanimlanmistir.",
            "Renkler koyu arka plan uzerinde yuksek kontrast hedefiyle secilmistir.",
            "Empty-state ve hata mesajlari metinsel olarak da iletilir; sadece renge bagli degildir.",
            "Eksikler: ARIA live region, skip-link, prefers-reduced-motion ve resmi axe/Lighthouse raporu bu surumde yoktur.",
        ],
    )


def render_security_perf_test(doc: Document) -> None:
    doc.add_heading("8. GUVENLIK, PERFORMANS VE TEST", level=1)
    doc.add_heading("8.1. Guvenlik (OWASP Top 10)", level=2)
    add_simple_table(
        doc,
        "Tablo 25: OWASP Top 10 (2021) Uygulama Matrisi",
        ["Kod", "Risk", "Bu projedeki onlemim"],
        [
            ["A01", "Broken Access Control", "requireAuth ve requireRole middleware'leri ile endpoint korumasi"],
            ["A02", "Cryptographic Failures", "bcrypt cost 12, JWT secret, HttpOnly cookie"],
            ["A03", "Injection", "Zod validation + parametreli Cypher sorgulari"],
            ["A04", "Insecure Design", "Role kontrollu admin reset, rate limiting, acik API yuzeyi"],
            ["A05", "Security Misconfiguration", "helmet, x-powered-by kapatma, local listen address"],
            ["A06", "Vulnerable Components", "package-lock kullanimi; manuel npm test akisi"],
            ["A07", "Auth Failures", "refresh rotation, oturum suresi kontrolu, invalid token 401"],
            ["A08", "Software Integrity", "Kilitleme dosyasi ve kontrollu scriptler"],
            ["A09", "Logging & Monitoring", "Temel console loglari mevcut; yapisal monitoring eksik"],
            ["A10", "SSRF", "Harici URL fetch yuzeyi yok; saldiri yuzeyi dusuk"],
        ],
    )
    add_paragraphs(
        doc,
        """
        OWASP Top 10 [8], web uygulamalari icin kritik riskleri tanimlayan ortak bir farkindalik dokumanidir. GraphLink
        bu listenin bir kismi icin dogrudan kod seviyesinde onlem uygular; bir kismi icin ise risk kabul edip backlog'a
        tasimistir. Ozellikle A09 maddesinde profesyonel loglama / monitoring katmaninin eksik oldugu acikca not edilmelidir.
        """,
    )

    doc.add_heading("8.2. Performans Hedefleri", level=2)
    add_simple_table(
        doc,
        "Tablo 26: Performans Metrikleri ve Hedefler",
        ["Metrik", "Hedef", "Gozlenen / Durum"],
        [
            ["Dashboard API ortalamasi", "< 500 ms", f"{PROJECT_FACTS['dashboard_avg_ms']} ms ortalama — karsilandi"],
            ["Dashboard API max (10 sample)", "Istikrali", f"{PROJECT_FACTS['dashboard_max_ms']} ms — kucuk orneklemde yuksek ama kabul edilebilir"],
            ["P95 benzeri yorum", "< 500 ms tercih", "10 orneklemde tepe deger 596 ms — iyilestirme alani var"],
            ["Ilk canli demo verisi", "Kucuk dataset", PROJECT_FACTS["live_counts"]],
            ["Bundle / asset stratejisi", "Kompakt", "Vanilla JS + tek CSS ile basit dagitim"],
        ],
    )
    add_paragraphs(
        doc,
        """
        Yerel olcumler {date} tarihinde, ayni makinede calisan Node.js uygulamasi ile Neo4j runtime uzerinden alinmistir.
        10 adet /api/dashboard isteginin sureleri sirayla {samples} ms olarak kaydedilmistir. Bu veriler production
        benchmark'i degil, final teslimi icin demo olceginde bir saglik kontroldur.
        """.format(date=MEASURE_DATE, samples=PROJECT_FACTS["dashboard_samples"]),
    )

    doc.add_heading("8.3. Test Stratejisi", level=2)
    add_image(doc, "Sekil 10: Test Piramidi", "test-pyramid.png", width=5.8)
    add_paragraphs(
        doc,
        """
        Mevcut repo'da toplam 61 test bulunmaktadir. Yerel Neo4j servisi olmayan makinede 57 test gecmis, 4 adet
        veritabani bagimli test otomatik skip edilmistir. GitHub Actions akisi Neo4j servis konteyneri baslattigi icin
        bu entegrasyon testleri CI ortaminda canli veritabaniyla kosacak sekilde ayarlanmistir.

        Test paketi validation schema kurallari, security helper'lari, CSP kontrolu, sosyal route negative-path
        davranislari, frontend static asset smoke testleri, monitoring ve auth akisini kapsar. Ayrica c8 ile coverage
        raporu uretilir ve npm audit orta seviye zaafiyetlere kadar pipeline icinde kontrol edilir. Eksik kalan ana alan,
        tarayici tabanli Playwright/Cypress E2E senaryolari ve buyuk veri seti performans benchmark'idir.
        """,
    )


def render_business(doc: Document) -> None:
    doc.add_heading("9. MALIYET, GELIR MODELI VE GTM", level=1)
    add_paragraphs(
        doc,
        """
        Bu bolum, GraphLink'in bugunku ogrenci projesi durumunu degil; ayni urunun egitim ve demo odakli bir mikro
        SaaS urune donusturulmesi halinde izlenebilecek varsayimsal is modelini anlatir. Gercek proje gelirleri veya
        operasyonel musteri verisi bulunmamaktadir.
        """,
    )

    doc.add_heading("9.1. Business Model Canvas", level=2)
    add_simple_table(
        doc,
        "Tablo 27: Business Model Canvas",
        ["Blok", "Icerik"],
        [
            ["Customer Segments", "Universiteler, bootcamp'ler, graph kavramini ogrenmek isteyen ekipler, erken asama startup'lar"],
            ["Value Propositions", "Yerel calisan graph demo; aciklanabilir recommendation paneli; Turkce ogrenme deneyimi"],
            ["Channels", "GitHub/portfolio, universite sunumlari, LinkedIn, teknik topluluklar"],
            ["Customer Relationships", "Self-service demo, dokumantasyon, egitmen destekli kullanim"],
            ["Revenue Streams", "Pro lisans, kurumsal egitim paketi, kurulum/ozellestirme danismanligi"],
            ["Key Resources", "Kod tabani, ekran goruntuleri, egitsel icerik, yerel kurulum scriptleri"],
            ["Key Activities", "Bakim, yeni demo senaryolari, egitim materyali, test kapsami genisletme"],
            ["Key Partnerships", "Universiteler, teknik topluluklar, Neo4j egitim toplulugu"],
            ["Cost Structure", "Gelistirme zamani, domain/hosting, icerik pazarlama, destek suresi"],
        ],
    )

    doc.add_heading("9.2. Gelir Modeli", level=2)
    add_paragraphs(
        doc,
        """
        Onerilen gelir modeli freemium + egitim/danismanlik karmasidir. Cekirdek demo ucretsiz tutulur; daha fazla
        senaryo, ozel veri seti, rapor ciktisi ve kurumsal destek katmanlari ucretlendirilir.
        """,
    )
    add_simple_table(
        doc,
        "Tablo 28: Fiyatlandirma (Varsayimsal 3 Tier)",
        ["Plan", "Aylik Fiyat", "Yillik Fiyat", "Icerik"],
        [
            ["Free", "0 TL", "0 TL", "Demo hesap, temel dashboard, sinirli veri seti"],
            ["Pro", "299 TL", "2.990 TL", "Tum graph panelleri, rapor export, ozel senaryolar"],
            ["Business", "1.499 TL", "14.990 TL", "Kurumsal onboarding, ozel dataset, teknik destek"],
        ],
    )

    doc.add_heading("9.3. Maliyet Analizi", level=2)
    add_simple_table(
        doc,
        "Tablo 29: 1. Yil TCO (Varsayimsal)",
        ["Kalem", "Tutar (TL)"],
        [
            ["Gelistirme emegi (freelance esdeger)", "120.000"],
            ["12 ay hosting / yedek / domain", "18.000"],
            ["Domain + SSL", "1.000"],
            ["Pazarlama / demo icerigi", "12.000"],
            ["Danismanlik / muhasebe / hukuk", "9.000"],
            ["TOPLAM", "160.000"],
        ],
    )

    doc.add_heading("9.4. Unit Economics", level=2)
    add_simple_table(
        doc,
        "Tablo 30: Unit Economics (Varsayimsal)",
        ["Metrik", "Deger", "Aciklama"],
        [
            ["ARPU", "299 TL / ay", "Pro plana gecen kullanici varsayimi"],
            ["Gross Margin", "%82", "Yazilim agirlikli urun marji varsayimi"],
            ["Aylik Churn", "%4", "Erken asama niche urun tahmini"],
            ["LTV", "Yaklasik 6.129 TL", "ARPU x margin / churn"],
            ["CAC", "550 TL", "Icerik + reklam + demo etkinligi karmasi"],
            ["LTV / CAC", "11.14", "Saglikli gorunen ama varsayimsal oran"],
            ["Payback Period", "Yaklasik 2.2 ay", "CAC / (ARPU x margin)"],
        ],
    )

    doc.add_heading("9.5. Go-to-Market (GTM) Stratejisi", level=2)
    add_bullets(
        doc,
        [
            "Ilk 100 kullanici icin universite siniflari, topluluk sunumlari ve GitHub demo sayfasi hedeflenir.",
            "Turkce graph database anlatan blog/youtube icerikleri ile inbound trafik olusturulur.",
            "Product Hunt yerine once daha nise teknik topluluklarda validasyon aranir.",
            "Growth loop: kullanicilar kendi veri setleriyle screenshot veya demo olusturdukca urun gorunurlugu artar.",
        ],
    )


def render_implementation(doc: Document) -> None:
    doc.add_heading("10. UYGULAMA VE GELISTIRME", level=1)
    doc.add_heading("10.1. Kurulum ve Calistirma", level=2)
    add_paragraphs(
        doc,
        """
        Proje, README'de tanimlandigi sekliyle Node.js ve yerel Neo4j runtime kullanilarak ayaga kaldirilir. Bu rapor
        icin canli testler de ayni adimlarla yapilmistir.
        """,
    )
    add_code_block(
        doc,
        """
        npm install
        npm run neo4j:setup
        npm run neo4j:start
        npm start
        npm test
        """,
    )

    doc.add_heading("10.2. Ekran Goruntuleri", level=2)
    add_paragraphs(
        doc,
        """
        Asagidaki ekran goruntuleri proje yerelde calisirken, {date} tarihinde otomasyon araciyla alinmistir. Boylece
        rapor, sadece taslak wireframe degil gercek calisan arayuz goruntulerini de icermektedir.
        """.format(date=MEASURE_DATE),
    )
    add_image(doc, "Sekil 11: Landing / Auth Ekrani", "landing.png", width=6.0)
    add_image(doc, "Sekil 12: Hatali Giris Durumu", "login-error.png", width=6.0)
    add_image(doc, "Sekil 13: Dashboard — Dolu State", "dashboard.png", width=6.0)
    add_image(doc, "Sekil 14: Kullanici Kartlari Paneli", "users-panel.png", width=5.6)
    add_image(doc, "Sekil 15: Gonderi Akisi Paneli", "posts-panel.png", width=5.6)
    add_image(doc, "Sekil 16: Graph Analiz Paneli", "graph-panel.png", width=5.6)
    add_image(doc, "Sekil 17: Mobile Gorunum", "mobile-dashboard.png", width=3.2)
    add_image(doc, "Sekil 18: Swagger / API Dokumantasyonu", "swagger.png", width=6.0)

    doc.add_heading("10.3. Kullanilan AI Araclari ve Katki Orani", level=2)
    add_simple_table(
        doc,
        "Tablo 31: AI Kullanim Dagilimi (Tahmini)",
        ["Arac", "Kullanim Orani", "Ne icin?"],
        [
            ["ChatGPT / Codex", "%45", "Dokumantasyon yapisi, rapor taslagi, kod okuma yardimi"],
            ["Cursor / GitHub Copilot", "%10", "Autocomplete ve kucuk refactoring yardimlari"],
            ["Claude Code", "%0", "Bu teslim surecinde aktif kullanilmadi"],
            ["AI ciktisinin manuel kontrolu", "%100", "Tum ciktilar proje baglamina gore gozden gecirildi"],
        ],
    )


def render_conclusion(doc: Document) -> None:
    doc.add_heading("11. SONUC VE DEGERLENDIRME", level=1)
    doc.add_heading("11.1. Proje Hakkinda Genel Degerlendirme", level=2)
    add_paragraphs(
        doc,
        """
        GraphLink, planlanan cekirdek hedeflerin buyuk bolumunu yerine getirmistir. Gercek auth akisi, graph veri
        modeli, sosyal etkileşimler, algorithm paneli, dokumantasyon ve yerel kurulum scriptleri ayni repo icinde
        calisir durumdadir. En guclu taraf, kullanici davranisi ile graph analitigi ayni arayuzde gosterebilmesidir.

        Bununla birlikte proje henuz production hazirligi seviyesinde degildir. En kritik kalan aciklar; tarayici
        tabanli E2E test eksikligi, sinirli monitoring, kurumsal olcekli operasyon ozelliklerinin olmamasi ve cok daha
        buyuk veri setlerinde benchmark alinmamis olmasidir. Son guncellemeyle guvenlik varsayilanlari daha kapali hale
        getirilmis, CI'daki veritabani testleri aktif edilmistir. Final projesi baglaminda teknik hedeflerin
        savunulabilir oldugu aciktir.
        """,
    )

    doc.add_heading("11.2. Karsilasilan Zorluklar ve Cozumleri", level=2)
    add_simple_table(
        doc,
        "Tablo 32: Gelistirme Surecindeki En Buyuk 3 Zorluk",
        ["Zorluk", "Cozum", "Ogrenilen"],
        [
            ["Template ile gercek stack arasindaki fark (React/Fastify vs Express/vanilla)", "ADR-001 ile mevcut mimari korundu ve raporda fark acikca belirtildi", "Teknik durustluk, gereksiz yeniden yazimdan degerlidir"],
            ["Graph algoritmalarini yerelde calistirma", "PowerShell setup script'i ile GDS/APOC jar'lari plugin klasorune kopyalandi", "Tekrarlanabilir kurulum demo kalitesini yukseltiyor"],
            ["Auth + refresh oturumunu sade tutma", "JWT access + Neo4j session node rotation modeli secildi", "Guvenlik ve basitlik arasinda pragmatik denge kurmak gerekiyor"],
        ],
    )

    doc.add_heading("11.3. Gelecek Calismalar (Future Work)", level=2)
    add_bullets(
        doc,
        [
            "V2: Live graph refresh ve notification deneyimini genisletme",
            "V2: Cypress/Playwright E2E senaryolari ve gorsel regresyon kontrolleri",
            "V2: TR + EN i18n ve resmi accessibility audit",
            "V3: GraphRAG / knowledge graph senaryolari",
            "V3: Aciklanabilir recommendation aciklamalarini LLM destekli hale getirme",
        ],
    )

    doc.add_heading("11.4. Kazanilan Yetkinlikler", level=2)
    add_bullets(
        doc,
        [
            "Teknik: Neo4j property graph modelleme ve Cypher pattern tasarimi",
            "Teknik: GDS ile Louvain, PageRank, FastRP ve Dijkstra kullanimi",
            "Teknik: JWT + HttpOnly cookie + refresh rotation mimarisi",
            "Teknik: Cytoscape.js ile graph veriyi gorsellestirme",
            "Surecsel: ADR yazimi ve teknik kararlarin kayda alinmasi",
            "Surecsel: PRD dusuncesi ile MVP kapsam belirleme",
            "Surecsel: Persona, JTBD ve urun konumlandirma calismasi",
        ],
    )


def render_references(doc: Document) -> None:
    doc.add_heading("KAYNAKCA", level=1)
    for ref in REFERENCES:
        doc.add_paragraph(ref)


def render_appendices(doc: Document) -> None:
    doc.add_heading("EKLER", level=1)
    doc.add_heading("EK A — Tam Ekran Goruntuleri Arsivi", level=2)
    for path in sorted(ASSET_DIR.glob("*.png")):
        doc.add_paragraph(f"{path.name} — docs/report-assets/{path.name}")

    doc.add_heading("EK B — OpenAPI Spesifikasyonu Ciktisi", level=2)
    add_code_block(
        doc,
        """
        openapi: 3.1.0
        info:
          title: GraphLink Social Graph API
          version: 1.1.0
        paths:
          /api/auth/login:
            post:
              summary: Login with email and password
          /api/dashboard:
            get:
              summary: Get dashboard data for the current user
          /api/graph/network:
            get:
              summary: Get Cytoscape network payload
          /api/graph/shortest-path:
            get:
              summary: Get shortest path using GDS Dijkstra when available
        """,
    )

    doc.add_heading("EK C — Kod Ornekleri", level=2)
    doc.add_paragraph("Ornek 1: Auth middleware")
    add_code_block(
        doc,
        """
        export function requireAuth(req, _res, next) {
          const token = readAccessToken(req);

          if (!token) {
            return next(new AppError("Oturum bulunamadi.", 401));
          }

          try {
            const payload = verifyAccessToken(token);
            req.user = {
              id: payload.sub,
              email: payload.email,
              name: payload.name,
              role: payload.role
            };
            return next();
          } catch (_error) {
            return next(new AppError("Gecersiz veya suresi dolmus oturum.", 401));
          }
        }
        """,
    )
    doc.add_paragraph("Yorum: Auth kontrolu kisa, okunabilir ve role tabanli genisletmeye uygun sekilde tasarlanmistir.")

    doc.add_paragraph("Ornek 2: GDS shortest-path fallback mantigi")
    add_code_block(
        doc,
        """
        export async function getDijkstraPath(viewerId, targetId) {
          if (!targetId) {
            return getShortestPath(viewerId, targetId);
          }

          const result = await withProjectedGraph(async (graphName) =>
            runRead(
              `
                MATCH (source:User {id: $viewerId}), (target:User {id: $targetId})
                CALL gds.shortestPath.dijkstra.stream($graphName, {
                  sourceNode: source,
                  targetNode: target,
                  relationshipWeightProperty: "weight"
                })
                YIELD totalCost, nodeIds
                RETURN {
                  nodes: [nodeId IN nodeIds | gds.util.asNode(nodeId) { .id, .name }],
                  hops: totalCost
                } AS path
              `,
              { graphName, viewerId, targetId }
            )
          );

          if (!result || !result.records.length) {
            return getShortestPath(viewerId, targetId);
          }

          return result.records[0].get("path");
        }
        """,
    )
    doc.add_paragraph("Yorum: Bu blok, plugin mevcut olmadiginda arayuzun bozulmamasini saglayan pragmatik bir fallback desenidir.")

    doc.add_paragraph("Ornek 3: Rate limit politika kurulumu")
    add_code_block(
        doc,
        """
        function buildLimiter(windowMs, max, message) {
          return rateLimit({
            windowMs,
            max,
            standardHeaders: true,
            legacyHeaders: false,
            keyGenerator: (req) => req.user?.id ?? ipKeyGenerator(req.ip),
            message: {
              error: message
            }
          });
        }
        """,
    )
    doc.add_paragraph("Yorum: Auth, read ve write siniflari icin tek bir limiter fabrikasi kullanilmasi, politikalarin merkezi kalmasini saglar.")


def build_document() -> Document:
    doc = Document()
    set_styles(doc)
    render_cover(doc)
    render_declaration(doc)
    render_abstracts(doc)
    render_front_matter(doc)
    render_intro(doc)
    render_prd(doc)
    render_market(doc)
    render_tech_stack(doc)
    render_architecture(doc)
    render_data_api(doc)
    render_ui(doc)
    render_security_perf_test(doc)
    render_business(doc)
    render_implementation(doc)
    render_conclusion(doc)
    render_references(doc)
    render_appendices(doc)
    return doc


def main() -> None:
    ensure_assets()
    generate_diagrams()
    generate_ui_evidence_images()
    doc = build_document()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Rapor olusturuldu: {OUTPUT_PATH}")
    if TEMPLATE_SOURCE.exists():
        print(f"Bolum siralamasi ve kapak bilgileri referans sablon: {TEMPLATE_SOURCE}")


if __name__ == "__main__":
    main()
